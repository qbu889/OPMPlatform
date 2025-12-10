# utils/document_formatter.py
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def analyze_template_format(template_path):
    """
    分析模板文档的格式规范
    """
    try:
        doc = Document(template_path)
        format_rules = {
            'heading_styles': {},
            'paragraph_formats': [],
            'numbering_patterns': [],
            'font_rules': {}
        }

        # 分析标题样式
        for i, paragraph in enumerate(doc.paragraphs[:20]):  # 分析前20段
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.replace('Heading ', ''))
                format_rules['heading_styles'][level] = {
                    'font_size': paragraph.runs[0].font.size if paragraph.runs else None,
                    'bold': paragraph.runs[0].bold if paragraph.runs else False,
                    'alignment': paragraph.alignment
                }

        # 分析编号模式
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and (text[0].isdigit() or text.startswith(('1.', '2.', '3.', '①', '②'))):
                format_rules['numbering_patterns'].append({
                    'text_sample': text[:20],
                    'font_size': paragraph.runs[0].font.size if paragraph.runs else None,
                    'bold': paragraph.runs[0].bold if paragraph.runs else False
                })

        logger.info(f"模板格式分析完成: {format_rules}")
        return format_rules
    except Exception as e:
        logger.error(f"分析模板格式失败: {str(e)}")
        return {}


def apply_format_to_document(source_doc_path, template_format, output_path):
    """
    将模板格式应用到源文档
    """
    try:
        source_doc = Document(source_doc_path)
        formatted_doc = Document()

        # 复制模板的基础样式
        template_doc = Document('templates/功能需求_Demo模板.docx')
        formatted_doc.styles = template_doc.styles

        # 处理每个段落
        for para in source_doc.paragraphs:
            new_para = formatted_doc.add_paragraph()
            new_para.text = para.text

            # 复制段落的基本内容
            if para.runs:
                new_para.clear()  # 清除默认文本
                for run in para.runs:
                    new_run = new_para.add_run(run.text)

                    # 应用基础格式（如果模板中有相关信息）
                    if 'heading_styles' in template_format and para.style.name.startswith('Heading'):
                        # 应用标题格式
                        level = int(para.style.name.replace('Heading ', '')) if para.style.name.replace('Heading ',
                                                                                                        '').isdigit() else 1
                        if level in template_format['heading_styles']:
                            style_info = template_format['heading_styles'][level]
                            if style_info['font_size']:
                                new_run.font.size = style_info['font_size']
                            if style_info['bold']:
                                new_run.bold = True
                    elif para.text.strip() and (para.text.strip()[0].isdigit() or
                                                any(pattern in para.text for pattern in ['1.', '2.', '3.', '①', '②'])):
                        # 应用编号格式
                        new_run.bold = True
                        new_run.font.size = Pt(12)  # 默认编号字号
                    else:
                        # 应用普通段落格式
                        new_run.font.size = Pt(10.5)

            # 特殊处理：检测并格式化编号项
            format_numbered_items(new_para)

        # 保存格式化后的文档
        formatted_doc.save(output_path)
        logger.info(f"文档格式化完成，保存至: {output_path}")
        return True, ""
    except Exception as e:
        logger.error(f"应用格式失败: {str(e)}")
        return False, str(e)


def format_numbered_items(paragraph):
    """
    格式化编号项目
    """
    text = paragraph.text.strip()
    if not text:
        return

    # 检测常见的编号模式
    numbered_patterns = [
        r'^\d+\.',  # 1. 2. 3.
        r'^\d+\)',  # 1) 2) 3)
        r'^[①②③④⑤⑥⑦⑧⑨⑩]',  # ① ② ③
        r'^[一二三四五六七八九十]、',  # 一、二、三、
    ]

    is_numbered = any(text.startswith(tuple([pattern.rstrip('$^') for pattern in numbered_patterns])))

    if is_numbered:
        # 对编号项目应用加粗和适当字号
        for run in paragraph.runs:
            run.bold = True
            if not run.font.size:
                run.font.size = Pt(12)


# utils/document_formatter.py (增强版本)

def enhanced_apply_format_to_document(source_doc_path, template_format, output_path):
    """
    增强版文档格式化函数
    """
    try:
        source_doc = Document(source_doc_path)
        formatted_doc = Document()

        # 复制模板样式
        template_doc = Document('templates/功能需求_Demo模板.docx')
        formatted_doc.styles = template_doc.styles

        # 记录处理统计信息
        stats = {
            'headings': 0,
            'numbered_items': 0,
            'paragraphs': 0
        }

        # 处理每个段落
        for i, para in enumerate(source_doc.paragraphs):
            new_para = formatted_doc.add_paragraph()

            # 检测段落类型并应用相应格式
            para_type = detect_paragraph_type(para.text)

            if para_type == 'heading':
                copy_and_format_heading(para, new_para, template_format)
                stats['headings'] += 1
            elif para_type == 'numbered':
                copy_and_format_numbered_item(para, new_para, template_format)
                stats['numbered_items'] += 1
            else:
                copy_and_format_paragraph(para, new_para, template_format)
                stats['paragraphs'] += 1

        # 保存文档
        formatted_doc.save(output_path)
        logger.info(f"文档格式化完成，统计信息: {stats}")
        return True, ""
    except Exception as e:
        logger.error(f"增强格式化失败: {str(e)}")
        return False, str(e)


def detect_paragraph_type(text):
    """
    检测段落类型
    """
    text = text.strip()
    if not text:
        return 'empty'

    # 检测标题（以数字开头且后面跟着顿号或点号）
    if text.startswith(('第', '附')) or \
            (len(text) > 1 and text[0].isdigit() and text[1] in ['.', '、', '章', '节']):
        return 'heading'

    # 检测编号项
    numbered_patterns = [
        r'^\d+(\.\d+)*',  # 1. 1.1 1.1.1
        r'^[①②③④⑤⑥⑦⑧⑨⑩]',
        r'^[一二三四五六七八九十]+[、.]',
        r'^\(\d+\)',
        r'^[A-Za-z]\.',
    ]

    for pattern in numbered_patterns:
        if text[:10].startswith(tuple([p.rstrip('*^$') for p in numbered_patterns])):
            return 'numbered'

    return 'normal'


# 修改copy_and_format_heading函数中的字体设置
def copy_and_format_heading(source_para, target_para, template_format):
    target_para.text = source_para.text
    level = 1  # 默认一级标题
    text = source_para.text.strip()
    if '章' in text:
        level = 1
    elif '节' in text or ('.' in text and text[0].isdigit()):
        level = 2
    else:
        level = 3

    for run in target_para.runs:
        run.bold = True
        run.font.size = Pt(16 - (level - 1) * 2)
        run.font.name = '宋体'  # 改为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 确保中文字体


def copy_and_format_numbered_item(source_para, target_para, template_format):
    """
    格式化编号项
    """
    target_para.text = source_para.text

    # 应用编号格式
    for run in target_para.runs:
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def copy_and_format_paragraph(source_para, target_para, template_format):
    """
    格式化普通段落
    """
    target_para.text = source_para.text

    # 应用普通段落格式
    for run in target_para.runs:
        run.bold = False
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
