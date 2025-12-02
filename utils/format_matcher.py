# utils/format_matcher.py
import logging
from docx import Document

logger = logging.getLogger(__name__)


def match_document_format(input_doc_path, standard_template_path):
    """
    匹配输入文档格式与标准模板格式
    """
    try:
        logger.info(f"开始匹配文档格式: {input_doc_path}")

        # 解析输入文档
        input_doc = Document(input_doc_path)
        standard_doc = Document(standard_template_path)

        # 提取输入文档结构
        input_structure = extract_document_structure(input_doc)
        standard_structure = extract_document_structure(standard_doc)

        # 比较结构匹配度
        match_result = compare_structures(input_structure, standard_structure)

        logger.info(f"文档格式匹配完成，匹配度: {match_result['match_rate']:.2f}")
        return match_result

    except Exception as e:
        logger.error(f"文档格式匹配过程中发生错误: {str(e)}")
        return {"match_rate": 0, "errors": [str(e)]}


def extract_document_structure(doc):
    """
    提取文档结构信息
    """
    structure = {
        "headings": [],
        "content_blocks": [],
        "field_patterns": []
    }

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            heading_info = {
                "level": int(para.style.name.replace('Heading ', '')),
                "text": para.text.strip(),
                "style": para.style.name
            }
            structure["headings"].append(heading_info)
        elif para.text.strip():
            # 检查是否包含特定字段模式
            text = para.text.strip()
            if any(keyword in text for keyword in ['功能描述：', '输入：', '输出：', '处理过程：']):
                structure["field_patterns"].append({
                    "pattern": detect_field_pattern(text),
                    "text": text
                })

    return structure


def detect_field_pattern(text):
    """
    检测字段模式
    """
    patterns = {
        "功能描述": "功能描述：" in text,
        "输入": "输入：" in text,
        "输出": "输出：" in text,
        "处理过程": "处理过程：" in text
    }
    return patterns


def compare_structures(input_struct, standard_struct):
    """
    比较两个文档结构
    """
    result = {
        "match_rate": 0.0,
        "errors": [],
        "details": {}
    }

    # 比较标题层级结构
    heading_match = compare_headings(input_struct["headings"], standard_struct["headings"])
    result["details"]["heading_match"] = heading_match

    # 比较字段模式
    field_match = compare_field_patterns(input_struct["field_patterns"], standard_struct["field_patterns"])
    result["details"]["field_match"] = field_match

    # 计算总体匹配度
    total_weight = 0.7 * heading_match + 0.3 * field_match
    result["match_rate"] = total_weight

    return result


def compare_headings(input_headings, standard_headings):
    """
    比较标题结构
    """
    if not standard_headings:
        return 1.0 if not input_headings else 0.0

    # 检查必要的层级是否存在
    required_levels = set(h["level"] for h in standard_headings)
    input_levels = set(h["level"] for h in input_headings)

    matched_levels = len(required_levels.intersection(input_levels))
    total_required = len(required_levels)

    return matched_levels / total_required if total_required > 0 else 1.0


def compare_field_patterns(input_patterns, standard_patterns):
    """
    比较字段模式
    """
    if not standard_patterns:
        return 1.0 if not input_patterns else 0.0

    required_fields = set()
    for pattern in standard_patterns:
        for field, present in pattern["pattern"].items():
            if present:
                required_fields.add(field)

    input_fields = set()
    for pattern in input_patterns:
        for field, present in pattern["pattern"].items():
            if present:
                input_fields.add(field)

    matched_fields = len(required_fields.intersection(input_fields))
    total_required = len(required_fields)

    return matched_fields / total_required if total_required > 0 else 1.0
