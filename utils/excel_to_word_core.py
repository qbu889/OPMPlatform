# utils/excel_to_word_core.py
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def read_excel_robust(file_path):
    """健壮读取Excel文件，自动识别表头"""
    try:
        df = pd.read_excel(file_path, header=0, engine='openpyxl')
        # 清理空行和空列
        df = df.dropna(how='all').dropna(axis=1, how='all')
        return df
    except Exception as e:
        raise Exception(f"读取Excel失败：{e}")


def excel_to_word(excel_path, word_path):
    """核心转换逻辑：Excel转Word"""
    df = read_excel_robust(excel_path)
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('COSMIC模块转换结果', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 遍历Excel行，生成Word内容（示例逻辑，可根据实际需求调整）
    for idx, row in df.iterrows():
        # 一级模块
        if pd.notna(row.get('一级模块')):
            doc.add_heading(row['一级模块'], level=1)
        # 二级模块
        if pd.notna(row.get('二级模块')):
            doc.add_heading(row['二级模块'], level=2)
        # 功能描述
        if pd.notna(row.get('功能描述')):
            para = doc.add_paragraph(row['功能描述'])
            para.font.size = Pt(11)

    doc.save(word_path)


def verify_consistency(excel_path, word_path):
    """校验Excel与Word内容一致性，返回校验结果+模块统计"""
    # 读取Excel数据
    df = read_excel_robust(excel_path)
    # 读取Word内容
    doc = Document(word_path)
    word_text = '\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    # 简单校验：核心字段是否存在
    core_fields = ['一级模块', '二级模块', '功能描述']
    verify_pass = True
    module_stats = []

    for field in core_fields:
        if field not in df.columns:
            verify_pass = False
            continue
        # 统计模块数据
        module_groups = df.groupby('一级模块').agg({
            '二级模块': 'nunique',
            '功能描述': 'count'
        }).reset_index()
        module_groups.columns = ['一级模块名称', '二级模块数量', '子过程数量']
        module_groups['CFP总和'] = 0  # 示例：根据实际逻辑计算CFP
        module_stats = module_groups.to_dict('records')

        # 校验每个模块是否在Word中存在
        for module in df[field].dropna().unique():
            if str(module) not in word_text:
                verify_pass = False

    return verify_pass, module_stats