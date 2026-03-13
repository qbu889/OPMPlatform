from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# 创建新的Word文档
doc = Document()

# 设置文档整体字体（解决中文显示问题）
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 添加标题
title = doc.add_heading('系统问题清单', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
title_run = title.runs[0]
title_run.font.size = Pt(18)  # 标题字号
title_run.font.bold = True    # 标题加粗

# 1. 基础功能异常
doc.add_heading('一、基础功能异常', level=1)
basic_problems = [
    '【延期申请-流转信息】延期申请、挂起申请留痕需补充派往对象（支持点击维护组查询对应人员）',
    '【任务详情】任务详情页显示“无子单”，但跳转至工单侧子单列表可查看到子单',
    '【作战室留痕】延期/挂起申请时，需新增留痕内容：“已通知xxx维护组进行审核；”（xxx维护组需显示为蓝色，且支持点击查看详情）',
    '【作战室留痕】调度指令存在错误'
]
for idx, problem in enumerate(basic_problems, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 2. 质检规则异常
doc.add_heading('二、质检规则异常', level=1)
quality_problems = [
    '无事件清除时间+无附件、故障短时恢复随意延期的质检项已停用，但仍触发该两项质检',
    '处理措施缺失关键字的质检项存在判定错误'
]
for idx, problem in enumerate(quality_problems, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 3. 手动建单相关问题
doc.add_heading('三、手动建单相关问题', level=1)
build_problems = [
    '【手动建单】事件流水号默认填充，需取消该默认填充逻辑',
    '【手动建单（已解决）】PC端新建任务时，时间组件若先选时分秒、不选日期，会自动跳转到1970年（测试环境已解决）',
    '【手动建单】PC端新建任务后，工单有事件描述，但任务侧无事件描述',
    '【手动建单】PC端新建任务后，省内派单级别显示异常（值为-1）'
]
for idx, problem in enumerate(build_problems, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 4. 人员调度与任务处理异常
doc.add_heading('四、人员调度与任务处理异常', level=1)
dispatch_problems = [
    '【人员调度决策-作战室留痕】汇聚（骨干）机房场景下，若静态表无升级人员，留痕需指派到区调度员（区调度员文字为蓝色、支持点击查询维护组人员）',
    '【人员决策】固定油机非局楼触发「油料补充」后，任务不应指派到人，需处于“待认领”状态',
    '【任务详情-改派-组内改派】生产环境调用组内改派接口返回空列表（测试环境正常），接口请求地址：http://10.44.225.27:8022/smart-module-src/monitor-intelligence-maintain/orgRole/getPersonChoiceList',
    '【任务认领】固定油机非局楼触发「油料补充」后，用户chenweiqiang_nok认领待认领任务时报错（code:500，msg:任务认领失败），接口请求地址：http://10.44.225.27:8021/smart-module-src/monitor-intelligence-maintain/taskInfo/taskAssignment'
]
for idx, problem in enumerate(dispatch_problems, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 5. 流转信息异常
doc.add_heading('五、流转信息异常', level=1)
flow_problems = [
    '【中间表】中间表导出后，需将“网管告警ID”左侧的“任务状态名称”字段名修改为“工单任务状态”',
    '【流转信息异常】T1移交T2时，派往对象内容错误（当前显示为“1”）',
    '【流转信息异常】T2确认受理时，操作人角色显示错误（当前：福州永泰农村网格虹信基站维护组2；预期：xxxx（T2）维护组）'
]
for idx, problem in enumerate(flow_problems, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 6. 应用功能优化（未发布生产）
doc.add_heading('六、应用功能优化（未发布生产）', level=1)
optimize_problems = [
    'PC端「全量列表」的任务类型筛选功能无效',
    'PC端点击任务编号无变色效果',
    'PC端无导出功能',
    'PC端待办列表无自动刷新功能'
]
for idx, problem in enumerate(optimize_problems, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 7. 任务回单异常
doc.add_heading('七、任务回单异常', level=1)
receipt_problem = [
    '【任务回单】任务已自动回单但界面未刷新，仍可提交质检，提交后报错提示需优化：\n'
    '  - 当前报错：提交质检异常，任务DH2026030301453不存在\n'
    '  - 预期提示：“提交质检失败，任务已闭环！”'
]
for idx, problem in enumerate(receipt_problem, 1):
    para = doc.add_paragraph()
    para.add_run(f'{idx}. ').bold = True
    para.add_run(problem)

# 保存文档
doc.save('系统问题清单.docx')
print('Word文档已生成：系统问题清单.docx')