# utils/demo_generator.py
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import logging
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def create_styles(doc):
    """创建符合要求的样式（统一宋体，按层级设置字号）"""
    # 字体大小映射（pt）：对应需求中的字号
    # 一级分类(Heading1)=三号(16pt)，二级=小三(15pt)，三级=四号(14pt)，四级=小四(12pt)，五级=小五(10.5pt)
    font_sizes = {1: 16, 2: 15, 3: 14, 4: 12, 5: 10.5}

    # 创建标题样式
    for i in range(1, 6):
        style_name = f'Custom Heading {i}'
        if style_name not in [s.name for s in doc.styles]:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles['Heading ' + str(i)]
            font = style.font
            font.size = Pt(font_sizes[i])
            font.name = '宋体'  # 统一使用宋体
            font.bold = True  # 标题默认加粗
            # 设置中文字体
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 配置主标题样式（5.功能需求：小二加粗宋体）
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.size = Pt(18)  # 小二对应18pt
    title_font.bold = True
    title_font.name = '宋体'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def generate_demo_doc(template_path, parsed_data, output_path):
    """生成符合格式要求的功能需求文档"""
    try:
        logger.info(f"开始生成文档: {output_path}")
        doc = Document()
        create_styles(doc)

        # 添加主标题（5.功能需求）
        title_para = doc.add_paragraph('5.功能需求')
        title_para.style = doc.styles['Title']

        # 添加"按照FPA"（加粗宋体）
        fpa_para = doc.add_paragraph("按照FPA")
        fpa_run = fpa_para.runs[0]
        fpa_run.bold = True
        fpa_run.font.name = '宋体'
        fpa_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 跟踪层级和计数器（保持原逻辑）
        current_levels = {'一级分类': None, '二级分类': None, '三级分类': None, '功能点名称': None}
        counters = {'一级': 0, '二级': 0, '三级': 0, '功能点': 0, '计数项': 0}

        for item in parsed_data:
            # 一级分类处理（Heading1：三号宋体）
            if item.get('一级分类') and item['一级分类'] != current_levels['一级分类']:
                counters['一级'] += 1
                counters['二级'] = counters['三级'] = counters['功能点'] = counters['计数项'] = 0
                current_levels['一级分类'] = item['一级分类']
                heading_text = f"{counters['一级']} {item['一级分类']}"
                para = doc.add_paragraph(heading_text)
                para.style = doc.styles['Custom Heading 1']

            # 二级分类处理（Heading2：小三宋体）
            if item.get('二级分类') and item['二级分类'] != current_levels['二级分类']:
                counters['二级'] += 1
                counters['三级'] = counters['功能点'] = counters['计数项'] = 0
                current_levels['二级分类'] = item['二级分类']
                heading_text = f"{counters['一级']}.{counters['二级']} {item['二级分类']}"
                para = doc.add_paragraph(heading_text)
                para.style = doc.styles['Custom Heading 2']

            # 三级分类处理（Heading3：四号宋体）
            if item.get('三级分类') and item['三级分类'] != current_levels['三级分类']:
                counters['三级'] += 1
                counters['功能点'] = counters['计数项'] = 0
                current_levels['三级分类'] = item['三级分类']
                heading_text = f"{counters['一级']}.{counters['二级']}.{counters['三级']} {item['三级分类']}"
                para = doc.add_paragraph(heading_text)
                para.style = doc.styles['Custom Heading 3']

            # 功能点名称处理（Heading4：小四宋体）
            if item.get('功能点名称') and item['功能点名称'] != current_levels['功能点名称']:
                counters['功能点'] += 1
                counters['计数项'] = 0
                current_levels['功能点名称'] = item['功能点名称']
                heading_text = f"{counters['一级']}.{counters['二级']}.{counters['三级']}.{counters['功能点']} {item['功能点名称']}"
                para = doc.add_paragraph(heading_text)
                para.style = doc.styles['Custom Heading 4']

            # 功能点计数项处理（剔除ILF/EIF，作为Heading5）
            if item.get('功能点计数项'):
                count_item = item['功能点计数项']
                # 剔除包含ILF或EIF的项
                if 'ILF' in count_item or 'EIF' in count_item:
                    logger.debug(f"已剔除ILF/EIF项: {count_item}")
                    continue
                # 保留项作为第五级标题
                counters['计数项'] += 1
                heading_text = f"{counters['一级']}.{counters['二级']}.{counters['三级']}.{counters['功能点']}.{counters['计数项']} {count_item}"
                para = doc.add_paragraph(heading_text)
                para.style = doc.styles['Custom Heading 5']

                # 添加详细信息（按格式要求设置标签样式）
                # 1. 场景说明：加粗
                if item.get('场景说明'):
                    para = doc.add_paragraph()
                    tag_run = para.add_run("场景说明：")
                    tag_run.bold = True
                    tag_run.font.name = '宋体'
                    tag_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    content_run = para.add_run(item['场景说明'])
                    content_run.font.name = '宋体'
                    content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 2. 功能描述：四号加粗
                if item.get('功能描述'):
                    para = doc.add_paragraph()
                    tag_run = para.add_run("功能描述：")
                    tag_run.font.size = Pt(14)  # 四号
                    tag_run.bold = True
                    tag_run.font.name = '宋体'
                    tag_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    content_run = para.add_run(item['功能描述'])
                    content_run.font.size = Pt(12)  # 内容用小四
                    content_run.font.name = '宋体'

                # 3. 系统界面：四号加粗，默认无
                para = doc.add_paragraph()
                tag_run = para.add_run("系统界面：")
                tag_run.font.size = Pt(14)
                tag_run.bold = True
                tag_run.font.name = '宋体'
                tag_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                content = item.get('系统界面', '无')  # 默认无
                content_run = para.add_run(content)
                content_run.font.size = Pt(12)
                content_run.font.name = '宋体'

                # 4. 输入：四号加粗
                if item.get('输入'):
                    para = doc.add_paragraph()
                    tag_run = para.add_run("输入：")
                    tag_run.font.size = Pt(14)
                    tag_run.bold = True
                    tag_run.font.name = '宋体'
                    content_run = para.add_run(item['输入'])
                    content_run.font.size = Pt(12)
                    content_run.font.name = '宋体'

                # 5. 输出：四号加粗
                if item.get('输出'):
                    para = doc.add_paragraph()
                    tag_run = para.add_run("输出：")
                    tag_run.font.size = Pt(14)
                    tag_run.bold = True
                    tag_run.font.name = '宋体'
                    content_run = para.add_run(item['输出'])
                    content_run.font.size = Pt(12)
                    content_run.font.name = '宋体'

                # 6. 处理过程：四号加粗
                if item.get('处理过程'):
                    para = doc.add_paragraph()
                    tag_run = para.add_run("处理过程：")
                    tag_run.font.size = Pt(14)
                    tag_run.bold = True
                    tag_run.font.name = '宋体'
                    content_run = para.add_run(item['处理过程'])
                    content_run.font.size = Pt(12)
                    content_run.font.name = '宋体'

                # 7. 涉及数据文件（FTR/RET）：四号加粗
                if item.get('涉及数据文件'):
                    para = doc.add_paragraph()
                    tag_run = para.add_run("本事务功能涉及到的数据文件（即FTR/RET）：")
                    tag_run.font.size = Pt(14)
                    tag_run.bold = True
                    tag_run.font.name = '宋体'
                    content_run = para.add_run(item['涉及数据文件'])
                    content_run.font.size = Pt(12)
                    content_run.font.name = '宋体'

                doc.add_paragraph("")  # 空行分隔

        doc.save(output_path)
        logger.info(f"文档生成完成: {output_path}")
        return True, None
    except Exception as e:
        error_msg = f"文档生成失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return False, error_msg
