#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 转 Excel 工具模块
解析 Word 技术规范书，自动生成软件资产清单 Excel 文件
"""
import os
import re
import logging
import pandas as pd
from typing import Dict, List, Optional, Tuple
from docx import Document
from pathlib import Path

logger = logging.getLogger(__name__)


class WordDocumentParser:
    """Word 文档解析器 - 用于解析技术规范书的层级结构"""
    
    def __init__(self):
        """初始化解析器"""
        self.data_rows = []  # 存储所有提取的数据行
        self.sequence_number = 0  # 序号计数器
        self.current_level1 = None  # 一级分类：如"监控管理应用" (Heading 5)
        self.current_level2 = None  # 二级分类：如"分级调度管理" (Heading 6)
        self.current_module = None  # 功能模块：如"分级调度派发" (Heading 7)

        
    def parse(self, filepath: str) -> List[Dict]:
        """
        解析 Word 文档
        
        Args:
            filepath: Word 文件路径
            
        Returns:
            包含所有功能点的列表，每个元素为字典：
            {
                '一级分类': str,
                '二级分类': str,
                '功能模块': str,
                '序号': int,
                '功能点名称': str,
                '功能点描述': str
            }
        """
        logger.info(f"[WORD_PARSER] 开始解析文档：{filepath}")
        
        try:
            doc = Document(filepath)
            logger.info(f"[WORD_PARSER] 文档加载成功，段落数：{len(doc.paragraphs)}")
            
            # 重置状态
            self.reset_state()
            
            # 遍历所有段落
            for para in doc.paragraphs:
                self._process_paragraph(para)
            
            logger.info(f"[WORD_PARSER] 解析完成，共提取 {len(self.data_rows)} 条功能点")
            return self.data_rows
            
        except Exception as e:
            logger.error(f"[WORD_PARSER] 解析失败：{e}", exc_info=True)
            raise Exception(f"Word 文档解析失败：{str(e)}")
    
    def reset_state(self):
        """重置解析状态"""
        self.current_level1 = None
        self.current_level2 = None
        self.current_module = None
        self.data_rows = []
        self.sequence_number = 0
    
    def _process_paragraph(self, paragraph):
        """
        处理单个段落
        
        Args:
            paragraph: docx 段落对象
        """
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else 'Normal'
        
        if not text:
            return
        
        # 根据标题样式识别层级
        if style == 'Heading 5':  # 一级分类
            self.current_level1 = text
            self.current_level2 = None
            self.current_module = None
            logger.debug(f"[WORD_PARSER] 一级分类：{text}")
        elif style == 'Heading 6':  # 二级分类
            self.current_level2 = text
            self.current_module = None
            logger.debug(f"[WORD_PARSER] 二级分类：{text}")
        elif style == 'Heading 7':  # 功能模块
            self.current_module = text
            logger.debug(f"[WORD_PARSER] 功能模块：{text}")
        elif style == 'Heading 8':  # 功能点名称
            # 提取功能点
            self._extract_function_point(text, paragraph)
        elif style in ['Normal', 'List Paragraph', '正文＋缩进 2 字符']:
            # 非标题段落，可能是功能点描述的延续
            # 检查是否属于某个功能点的描述
            if self.data_rows and self.current_module:
                # 将文本追加到最后一个功能点的描述中
                last_row = self.data_rows[-1]
                if last_row.get('功能点描述'):
                    last_row['功能点描述'] += '\n' + text
                else:
                    last_row['功能点描述'] = text
    
    def _parse_hierarchy_number(self, text: str) -> Optional[Dict]:
        """
        解析层级编号
        
        Args:
            text: 段落文本
            
        Returns:
            匹配到的层级信息：
            {
                'number': str,  # 完整编号，如"3.1.1.2.1.1.1.1"
                'title': str,   # 标题文本
                'level': int    # 层级深度（编号段数）
            }
            或者 None（如果不匹配）
        """
        # 匹配模式：数字开头的编号 + 空格或冒号 + 标题
        # 例如："3.1.1.2.1.1.1.1 派发智能体自动修复"
        pattern = r'^(\d+(?:\.\d+)*)[\s、](.+)$'
        
        match = re.match(pattern, text)
        if match:
            number = match.group(1)
            title = match.group(2).strip()
            
            # 计算层级深度
            levels = number.split('.')
            level = len(levels)
            
            # 只处理 4 级及以上的编号（一级分类至少是 4 级编号）
            if level >= 4:
                return {
                    'number': number,
                    'title': title,
                    'level': level
                }
        
        return None
    
    def _extract_function_point(self, title: str, paragraph):
        """
        提取功能点
        
        Args:
            title: 功能点名称
            paragraph: 段落对象
        """
        # 增加序号
        self.sequence_number += 1
        
        # 提取功能点描述
        description = self._extract_description(paragraph)
        
        # 创建数据行
        row = {
            '一级分类': self.current_level1 or '',
            '二级分类': self.current_level2 or '',
            '功能模块': self.current_module or '',
            '序号': self.sequence_number,
            '功能点名称': title,
            '功能点描述': description
        }
        
        self.data_rows.append(row)
        logger.debug(f"[WORD_PARSER] 提取功能点 #{self.sequence_number}: {title}")
    
    def _extract_description(self, paragraph) -> str:
        """
        提取功能点描述（纯文本，不包含图片）
        
        Args:
            paragraph: docx 段落对象
            
        Returns:
            功能点描述文本（保留原始格式和换行）
        """
        # 获取段落文本
        text = paragraph.text.strip()
        
        # 移除编号部分（如果有的话）
        # 例如："3.1.1.2.1.1.1.1 派发智能体自动修复" -> 只保留描述部分
        pattern = r'^\d+(?:\.\d+)*[\s、]'
        description = re.sub(pattern, '', text).strip()
        
        # 清理特殊字符和格式符号
        # 移除项目符号（如 、等）
        description = re.sub(r'^[•◦▪▸▹►‣⁃]\s*', '', description)
        
        # 保留换行符，只将多个连续空格替换为单个空格
        description = re.sub(r'[^\S\n]+', ' ', description).strip()
        
        return description


def generate_excel(data_rows: List[Dict], output_path: str):
    """
    生成 Excel 文件（带自动换行格式和美化样式）
    
    Args:
        data_rows: 数据行列表
        output_path: 输出文件路径
    """
    logger.info(f"[EXCEL_GEN] 开始生成 Excel 文件：{output_path}")
    
    try:
        # 创建 DataFrame
        df = pd.DataFrame(data_rows)
        
        # 确保列顺序正确 - 序号放在第一列
        column_order = ['序号', '一级分类', '二级分类', '功能模块', '功能点名称', '功能点描述']
        df = df[column_order]
        
        # 保存为 Excel 文件并使用 xlsxwriter 设置格式
        with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
            # 写入数据
            df.to_excel(writer, index=False, sheet_name='软件资产清单')
            
            # 获取 workbook 和 worksheet 对象
            workbook = writer.book
            worksheet = writer.sheets['软件资产清单']
            
            # ========== 创建各种格式 ==========
            
            # 表头格式：深蓝色背景，白色粗体
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'vcenter',
                'align': 'center',
                'bg_color': '#2E75B6',  # 深蓝色
                'font_color': '#FFFFFF',  # 白色
                'border': 1,
                'font_size': 11
            })
            
            # 序号列格式：居中显示（无边框）
            sequence_format = workbook.add_format({
                'align': 'center',
                'valign': 'vcenter',
                'font_size': 10
            })
            
            # 分类列格式：居中对齐（无边框）
            category_format = workbook.add_format({
                'align': 'center',
                'valign': 'vcenter',
                'font_size': 10,
                'text_wrap': True
            })
            
            # 功能点名称格式：居中对齐（无边框）
            function_name_format = workbook.add_format({
                'align': 'center',
                'valign': 'vcenter',
                'font_size': 10,
                'bold': True,
                'text_wrap': True
            })
            
            # 功能点描述格式：左对齐，自动换行（无边框）
            description_format = workbook.add_format({
                'align': 'left',
                'valign': 'top',
                'font_size': 10,
                'text_wrap': True
            })
            
            # 带边框的格式（用于有数据的行）
            sequence_border_format = workbook.add_format({
                'align': 'center',
                'valign': 'vcenter',
                'border': 1,
                'font_size': 10
            })
            
            category_border_format = workbook.add_format({
                'align': 'center',
                'valign': 'vcenter',
                'border': 1,
                'font_size': 10,
                'text_wrap': True
            })
            
            function_name_border_format = workbook.add_format({
                'align': 'center',
                'valign': 'vcenter',
                'border': 1,
                'font_size': 10,
                'bold': True,
                'text_wrap': True
            })
            
            description_border_format = workbook.add_format({
                'align': 'left',
                'valign': 'top',
                'border': 1,
                'font_size': 10,
                'text_wrap': True
            })
            
            # ========== 设置列宽 ==========
            
            # 序号列
            worksheet.set_column('A:A', 6, sequence_format)
            
            # 一级分类、二级分类、功能模块列
            worksheet.set_column('B:D', 15, category_format)
            
            # 功能点名称列
            worksheet.set_column('E:E', 22, function_name_format)
            
            # 功能点描述列（设置自动换行）
            worksheet.set_column('F:F', 45, description_format)
            
            # ========== 应用表头格式 ==========
            
            # 为表头（第 1 行）应用格式
            for col_num in range(len(df.columns)):
                worksheet.write(0, col_num, df.columns[col_num], header_format)
            
            # ========== 应用数据行格式（只为有数据的行添加边框） ==========
            
            # 为数据行应用格式（从第 2 行开始）
            for row_num in range(1, len(df) + 1):
                # 检查该行是否有数据
                row_has_data = any(df.iloc[row_num - 1].notna() & (df.iloc[row_num - 1] != ''))
                
                # 根据是否有数据选择格式
                seq_fmt = sequence_border_format if row_has_data else sequence_format
                cat_fmt = category_border_format if row_has_data else category_format
                name_fmt = function_name_border_format if row_has_data else function_name_format
                desc_fmt = description_border_format if row_has_data else description_format
                
                # 序号列
                worksheet.write(row_num, 0, df.iloc[row_num - 1]['序号'], seq_fmt)
                
                # 分类列（一级分类、二级分类、功能模块）
                for col_idx, col_name in enumerate(['一级分类', '二级分类', '功能模块']):
                    worksheet.write(row_num, col_idx + 1, df.iloc[row_num - 1][col_name], cat_fmt)
                
                # 功能点名称列
                worksheet.write(row_num, 4, df.iloc[row_num - 1]['功能点名称'], name_fmt)
                
                # 功能点描述列
                worksheet.write(row_num, 5, df.iloc[row_num - 1]['功能点描述'], desc_fmt)
            
            # ========== 设置行高（根据内容自动调整） ==========
            
            # 表头行高
            worksheet.set_row(0, 25)
            
            # 数据行行高（根据内容自动调整）
            for row_num in range(1, len(df) + 1):
                # 获取该行功能点描述的内容
                desc = str(df.iloc[row_num - 1]['功能点描述']) if pd.notna(df.iloc[row_num - 1]['功能点描述']) else ''
                
                # 根据内容长度计算行高
                if desc:
                    # 计算换行次数
                    line_count = desc.count('\n') + 1
                    # 根据行数和列宽计算行高（每行约 15 像素）
                    row_height = max(20, line_count * 15)
                    # 限制最大行高
                    row_height = min(row_height, 200)
                    worksheet.set_row(row_num, row_height)
                else:
                    # 没有内容的行设置较小的高度
                    worksheet.set_row(row_num, 15)
            
            # ========== 添加冻结窗格 ==========
            
            # 冻结第一行（表头）
            worksheet.freeze_panes(1, 0)
            
            # ========== 添加自动筛选 ==========
            
            # 为表头添加自动筛选
            worksheet.autofilter(0, 0, len(df), len(df.columns) - 1)
        
        logger.info(f"[EXCEL_GEN] Excel 文件生成成功：{output_path}")
        logger.info(f"[EXCEL_GEN] 共 {len(df)} 行数据")
        
        return True
        
    except Exception as e:
        logger.error(f"[EXCEL_GEN] Excel 生成失败：{e}", exc_info=True)
        raise Exception(f"Excel 文件生成失败：{str(e)}")


def parse_word_to_excel(input_path: str, output_path: str) -> bool:
    """
    Word 转 Excel 主函数
    
    Args:
        input_path: Word 文件输入路径
        output_path: Excel 文件输出路径
        
    Returns:
        是否成功
    """
    logger.info(f"[WORD_TO_EXCEL] 开始转换：{input_path} -> {output_path}")
    
    try:
        # 1. 解析 Word 文档
        parser = WordDocumentParser()
        data_rows = parser.parse(input_path)
        
        if not data_rows:
            logger.warning(f"[WORD_TO_EXCEL] 未提取到任何功能点")
            return False
        
        # 2. 生成 Excel 文件
        generate_excel(data_rows, output_path)
        
        logger.info(f"[WORD_TO_EXCEL] 转换完成，共提取 {len(data_rows)} 条功能点")
        return True
        
    except Exception as e:
        logger.error(f"[WORD_TO_EXCEL] 转换失败：{e}", exc_info=True)
        return False


def parse_word_to_excel_with_progress(input_path: str):
    """
    Word 转 Excel（带进度反馈）
    
    Args:
        input_path: Word 文件输入路径
        
    Returns:
        包含以下字段的字典：
        {
            'success': bool,
            'message': str,
            'excel_path': str (成功时),
            'total_functions': int (成功时)
        }
    """
    logger.info(f"[WORD_TO_EXCEL] 开始转换（带进度）：{input_path}")
    
    try:
        # 阶段 1: 解析 Word 文档 (0% - 60%)
        logger.info("[PROGRESS] 阶段 1: 解析 Word 文档...")
        parser = WordDocumentParser()
        data_rows = parser.parse(input_path)
        
        if not data_rows:
            return {
                'success': False,
                'message': '未提取到任何功能点，请检查文档格式是否正确'
            }
        
        logger.info(f"[PROGRESS] 解析完成，提取 {len(data_rows)} 条功能点")
        
        # 阶段 2: 生成 Excel 文件 (60% - 90%)
        logger.info("[PROGRESS] 阶段 2: 生成 Excel 文件...")
        
        # 创建临时输出路径
        import uuid
        output_filename = f"{Path(input_path).stem}_软件资产清单.xlsx"
        output_dir = os.path.dirname(input_path)
        output_path = os.path.join(output_dir, output_filename)
        
        generate_excel(data_rows, output_path)
        
        logger.info(f"[PROGRESS] Excel 生成完成")
        
        # 阶段 3: 完成 (90% - 100%)
        logger.info(f"[PROGRESS] 转换完成，共提取 {len(data_rows)} 条功能点")
        
        return {
            'success': True,
            'message': f'转换成功，共提取 {len(data_rows)} 条功能点',
            'excel_path': output_path,
            'total_functions': len(data_rows)
        }
        
    except Exception as e:
        logger.error(f"[WORD_TO_EXCEL] 转换失败：{e}", exc_info=True)
        return {
            'success': False,
            'message': f'转换失败：{str(e)}'
        }


# 便捷函数
def quick_convert(word_file: str, output_dir: str = None) -> Optional[str]:
    """
    快速转换 Word 为 Excel
    
    Args:
        word_file: Word 文件路径
        output_dir: 输出目录（可选，默认为 Word 文件所在目录）
        
    Returns:
        Excel 文件路径，如果失败则返回 None
    """
    try:
        # 确定输出目录
        if not output_dir:
            output_dir = os.path.dirname(word_file)
        
        # 生成输出文件名
        word_filename = Path(word_file).stem
        output_filename = f"{word_filename}_软件资产清单.xlsx"
        output_path = os.path.join(output_dir, output_filename)
        
        # 执行转换
        if parse_word_to_excel(word_file, output_path):
            return output_path
        else:
            return None
            
    except Exception as e:
        logger.error(f"[QUICK_CONVERT] 转换失败：{e}")
        return None


if __name__ == '__main__':
    # 测试代码
    test_word = '/Users/linziwang/PycharmProjects/wordToWord/docs/need_doc/监控综合应用系统 (2025-2026) 项目 - 技术规范书20250519.docx'
    test_output = '/Users/linziwang/PycharmProjects/wordToWord/docs/need_doc/监控综合应用系统 (2025-2026) 项目 - 软件资产解析清单.xlsx'
    
    if os.path.exists(test_word):
        result = quick_convert(test_word)
        if result:
            print(f"✅ 转换成功：{result}")
        else:
            print("❌ 转换失败")
    else:
        print(f"⚠️ 测试文件不存在：{test_word}")
