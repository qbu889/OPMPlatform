# utils/document_parser.py
from docx import Document

from app import logger

def validate_document_structure(doc):
    """验证文档是否包含必要结构"""
    # 检查文档是否具有预期的特征
    # 例如：特定标题、表格格式等
    pass
# utils/document_parser.py
def parse_source_doc(doc_path):
    logger.info(f"[DOC_PARSE] Starting document parsing: {doc_path}")
    try:
        doc = Document(doc_path)

        # 添加详细的文档结构信息日志
        logger.info(f"[DOC_PARSE] Paragraph count: {len(doc.paragraphs)}")
        logger.info(f"[DOC_PARSE] Table count: {len(doc.tables)}")

        # 记录前几个段落的内容（用于调试）
        for i, paragraph in enumerate(doc.paragraphs[:10]):
            logger.debug(f"[DOC_PARSE] Para {i}: {paragraph.text[:100]}...")

        # 记录表格信息（如果有的话）
        for i, table in enumerate(doc.tables[:3]):
            logger.debug(f"[DOC_PARSE] Table {i} dimensions: {len(table.rows)}x{len(table.columns)}")
            if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                logger.debug(f"[DOC_PARSE] Table {i} cell[0][0] content: {table.rows[0].cells[0].text[:50]}...")

        parsed_data = []

        # 记录文档的基本信息
        logger.info(f"[DOC_PARSE] Paragraph count: {len(doc.paragraphs)}")
        logger.info(f"[DOC_PARSE] Table count: {len(doc.tables)}")

        # 确保正确遍历所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:  # 跳过空段落
                continue

            # 检查段落样式和内容
            style_name = para.style.name
            logger.debug(f"[DOC_PARSE] Paragraph: '{text}' | Style: '{style_name}'")

            # 根据实际文档结构调整解析逻辑
            # 可能需要调整样式识别条件
            if style_name.startswith('Heading'):
                # 处理标题段落
                pass
            else:
                # 处理正文段落
                pass

        logger.info(f"[DOC_PARSE] Parsing completed, extracted {len(parsed_data)} items")
        return parsed_data, None
    except Exception as e:
        logger.error(f"[DOC_PARSE] Parsing error: {str(e)}", exc_info=True)
        return None, str(e)

