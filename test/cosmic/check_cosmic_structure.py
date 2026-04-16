from docx import Document
from pathlib import Path

base = Path('/')

# 对比两个文档的 H3 和 H4
cosmic_doc = Document(base / 'test/cosmic/附件2 2024年网络数据共享平台研发-功能点拆分表cosmic (5).docx')
req_doc = Document(base / 'test/cosmic/副本附件1-2024年网络数据共享平台研发-需求说明V1.0.0.docx')

print("=== COSMIC 功能点拆分表（前3组）===")
h3_h4_pairs = []
current_h3 = None
for para in cosmic_doc.paragraphs[:50]:
    if para.style.name == 'Heading 3':
        current_h3 = para.text
    elif para.style.name == 'Heading 4' and current_h3:
        h3_h4_pairs.append((current_h3, para.text))
        if len(h3_h4_pairs) >= 3:
            break

for h3, h4 in h3_h4_pairs:
    same = "✓ 相同" if h3 == h4 else "✗ 不同"
    print(f"H3: {h3}")
    print(f"H4: {h4} ({same})")
    print()

print("\n=== 需求说明文档（前3组）===")
h3_h4_pairs = []
current_h3 = None
for para in req_doc.paragraphs[:50]:
    if para.style.name == 'Heading 3':
        current_h3 = para.text
    elif para.style.name == 'Heading 4' and current_h3:
        h3_h4_pairs.append((current_h3, para.text))
        if len(h3_h4_pairs) >= 3:
            break

for h3, h4 in h3_h4_pairs:
    same = "✓ 相同" if h3 == h4 else "✗ 不同"
    print(f"H3: {h3}")
    print(f"H4: {h4} ({same})")
    print()
