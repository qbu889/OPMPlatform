#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试 .docx文件处理
"""
import sys
import os

# 添加项目根目录到路径
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
sys.path.insert(0, project_root)

from utils.document_processor import DocumentProcessor

def test_docx():
    """测试 .docx文件处理"""
    print("=" * 60)
    print("测试 .docx文件处理")
    print("=" * 60)
    
    processor = DocumentProcessor(upload_folder='test_uploads')
    
    # 创建测试 .docx文件
    docx_file = os.path.join(processor.upload_folder, 'test.docx')
    
    try:
        from docx import Document
        
        # 创建简单的 Word 文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试段落。')
        doc.add_paragraph('问题：什么是人工智能？')
        doc.add_paragraph('答案：人工智能是模拟人类智能的技术。')
        doc.save(docx_file)
        
        print(f"\n✅ 测试文档已创建：{docx_file}")
        
        # 测试处理文档
        result = processor.process_document(docx_file)
        
        if result['success']:
            print(f"✅ .docx文件处理成功")
            print(f"   文件类型：{result['filetype']}")
            print(f"   内容长度：{len(result['content'])} 字符")
            print(f"   内容预览：{result['content'][:100]}...")
        else:
            print(f"❌ .docx文件处理失败：{result.get('error')}")
            
    except Exception as e:
        print(f"❌ 测试异常：{e}")
        import traceback
        traceback.print_exc()
    finally:
        # 清理测试文件
        if os.path.exists(docx_file):
            os.remove(docx_file)
        # 清理测试目录
        if os.path.exists(processor.upload_folder):
            import shutil
            shutil.rmtree(processor.upload_folder, ignore_errors=True)
    
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)

if __name__ == '__main__':
    test_docx()
