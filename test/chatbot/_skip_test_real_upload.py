#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模拟真实文件上传测试
"""
import requests
import os
from docx import Document

# 创建测试文件
upload_folder = 'test_uploads'
os.makedirs(upload_folder, exist_ok=True)
test_file = os.path.join(upload_folder, '事件工单运维手册.docx')

doc = Document()
doc.add_heading('事件工单运维手册', 0)
doc.add_paragraph('这是测试内容。')
doc.save(test_file)

print(f"✅ 测试文件已创建：{test_file}")

# 上传文件
url = 'http://127.0.0.1:5000/chatbot/upload_document'
session_cookie = {
    'session': 'eyJjaGF0X3Nlc3Npb25faWQiOiJhYmQ0YzU4My0xNjk5LTRhMDktODMxNS1hNDE1ZWMwNzJhNWUifQ.aaqakg.RaBfScPmu8dYjvTirMW0595Pwu0'
}

try:
    with open(test_file, 'rb') as f:
        files = {'file': ('事件工单运维手册.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        
        print(f"\n📤 正在上传文件到：{url}")
        response = requests.post(url, files=files, cookies=session_cookie)
        
        print(f"\n响应状态码：{response.status_code}")
        print(f"响应内容：{response.json()}")
        
        if response.status_code == 200:
            print("\n✅ 上传成功！")
        else:
            print(f"\n❌ 上传失败：{response.json().get('error')}")
            
finally:
    # 清理测试文件
    if os.path.exists(test_file):
        os.remove(test_file)
        print(f"\n🗑️  测试文件已清理")
