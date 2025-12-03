# routes/document_routes.py
from flask import Blueprint, request, render_template, send_file, jsonify, current_app
import os
import logging
from docx import Document

from utils.demo_generator import generate_demo_doc
from utils.demo_validator import validate_demo_format
from utils.document_formatter import analyze_template_format, enhanced_apply_format_to_document
from utils.document_parser import parse_source_doc
from utils.format_matcher import match_document_format

# 创建蓝图
document_bp = Blueprint('document', __name__)
logger = logging.getLogger(__name__)


def inspect_document(doc_path):
    """检查文档基本信息"""
    try:
        doc = Document(doc_path)
        inspection_result = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'sample_paragraphs': [],
            'sample_tables': []
        }

        # 提取前几个段落样本
        for i, paragraph in enumerate(doc.paragraphs[:5]):
            inspection_result['sample_paragraphs'].append({
                'index': i,
                'text': paragraph.text[:100] + ('...' if len(paragraph.text) > 100 else '')
            })

        # 提取前几个表格样本
        for i, table in enumerate(doc.tables[:2]):
            table_info = {
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0,
                'sample_cells': []
            }
            if table.rows and len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                for j, cell in enumerate(table.rows[0].cells[:3]):
                    table_info['sample_cells'].append({
                        'col_index': j,
                        'text': cell.text[:50] + ('...' if len(cell.text) > 50 else '')
                    })
            inspection_result['sample_tables'].append(table_info)

        return inspection_result
    except Exception as e:
        logger.error(f"文档检查失败: {str(e)}")
        return None

@document_bp.route('/download-demo-template')
def download_demo_template():
    logger.info("开始处理下载Demo模板请求")
    # 修正路径为相对路径
    template_path = 'templates/功能需求_Demo模板.docx'
    if os.path.exists(template_path):
        logger.info(f"成功找到模板文件: {template_path}")
        return send_file(template_path, as_attachment=True, download_name='功能需求_Demo模板.docx')
    else:
        logger.error(f"模板文件不存在: {template_path}")
        return jsonify({'success': False, 'msg': '模板文件不存在'}), 404

@document_bp.route('/upload-demo', methods=['POST'])
def upload_demo():
    logger.info("开始处理Demo模板上传请求")
    if 'demo_file' not in request.files:
        logger.warning("未选择文件")
        return render_template('demo_upload.html', success=False, msg='未选择文件')

    demo_file = request.files['demo_file']
    if demo_file.filename == '':
        logger.warning("文件名为空")
        return render_template('demo_upload.html', success=False, msg='文件名不能为空')

    # 保存临时文件
    temp_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'temp_demo.docx')
    demo_file.save(temp_path)
    logger.info(f"临时文件已保存至: {temp_path}")

    # 验证Demo格式
    logger.info("开始验证Demo格式")
    validation_result, error_details = validate_demo_format(temp_path)
    logger.info(f"验证结果: {validation_result}, 错误数量: {len(error_details) if error_details else 0}")

    if not validation_result:
        logger.warning(f"Demo格式验证失败，错误详情: {error_details}")
        os.remove(temp_path)  # 删除无效文件
        return render_template('demo_upload.html', success=False, msg='Demo格式验证失败', error_details=error_details)

    # 验证通过：保存为正式模板
    os.rename(temp_path, current_app.config['DEMO_TEMPLATE_PATH'])
    logger.info(f"Demo模板验证通过，已保存为正式模板: {current_app.config['DEMO_TEMPLATE_PATH']}")
    return render_template('demo_upload.html', success=True, msg='Demo模板上传验证成功，可进行下一步文档转换')


@document_bp.route('/convert-page')
def convert_page():
    logger.info("访问转换页面")
    demo_exists = os.path.exists(current_app.config['DEMO_TEMPLATE_PATH'])
    logger.info(f"Demo模板存在状态: {demo_exists}")
    return render_template('convert_upload.html', demo_exists=demo_exists, convert_result=None)


@document_bp.route('/upload-and-convert', methods=['POST'])
def upload_and_convert():
    logger.info("开始处理文档转换请求")

    # 检查Demo模板是否存在
    if not os.path.exists(current_app.config['DEMO_TEMPLATE_PATH']):
        logger.warning("Demo模板不存在")
        # 添加格式化功能的提示
        return render_template('convert_upload.html', demo_exists=False,
                               convert_result=None, formatter_available=True)

    # 处理上传文件
    if 'convert_file' not in request.files:
        logger.warning("未选择待转换文件")
        return render_template('convert_upload.html', demo_exists=True,
                               convert_result={'success': False, 'msg': '未选择待转换文件'})

    convert_file = request.files['convert_file']
    if convert_file.filename == '':
        logger.warning("待转换文件名为空")
        return render_template('convert_upload.html', demo_exists=True,
                               convert_result={'success': False, 'msg': '文件名不能为空'})

    try:
        # 保存待转换文件
        source_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'source_doc.docx')
        convert_file.save(source_path)
        logger.info(f"待转换文件已保存至: {source_path}")

        # 文档预检
        inspection_result = inspect_document(source_path)
        if inspection_result:
            logger.info(
                f"文档检查结果: 段落数={inspection_result['paragraph_count']}, 表格数={inspection_result['table_count']}")

        # 步骤1：解析待转换文档，提取功能需求信息
        logger.info("开始解析待转换文档")
        parsed_data, parse_error = parse_source_doc(source_path)
        if parse_error:
            logger.error(f"文档解析失败: {parse_error}")
            return render_template('convert_upload.html', demo_exists=True,
                                   convert_result={'success': False, 'msg': f'文档解析失败：{parse_error}'})

        if not parsed_data or len(parsed_data) == 0:
            logger.warning("未解析到有效数据")
            # 提供更详细的错误信息帮助用户理解问题
            detailed_msg = ("文档中未找到有效的功能需求数据。请检查："
                           "1. 文档是否包含功能需求相关内容；"
                           "2. 文档格式是否符合要求；"
                           "3. 是否参考了标准模板格式")
            return render_template('convert_upload.html', demo_exists=True,
                                   convert_result={'success': False, 'msg': detailed_msg})

        logger.info(f"文档解析成功，共提取到 {len(parsed_data)} 条数据")

        # 步骤2：按Demo模板格式生成新文档
        output_doc_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'converted_demo.docx')
        logger.info("开始生成转换后的文档")

        generate_success, generate_error = generate_demo_doc(
            demo_template_path=current_app.config['DEMO_TEMPLATE_PATH'],
            parsed_data=parsed_data,
            output_path=output_doc_path
        )

        if not generate_success:
            logger.error(f"文档转换失败: {generate_error}")
            return render_template('convert_upload.html', demo_exists=True,
                                   convert_result={'success': False, 'msg': f'文档转换失败：{generate_error}'})

        logger.info(f"文档转换成功，输出文件路径: {output_doc_path}")

        # 转换成功，返回下载链接
        return render_template('convert_upload.html', demo_exists=True, convert_result={
            'success': True,
            'msg': f'文档已成功转换为Demo格式，共处理 {len(parsed_data)} 条数据',
            'download_url': '/download-converted-doc'
        })

    except Exception as e:
        logger.error(f"文档转换过程中发生未预期错误: {str(e)}", exc_info=True)
        return render_template('convert_upload.html', demo_exists=True,
                               convert_result={'success': False, 'msg': f'系统错误：{str(e)}'})
@document_bp.route('/download-converted-doc')
def download_converted_doc():
    logger.info("开始处理转换后文档下载请求")
    converted_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'converted_demo.docx')
    if not os.path.exists(converted_path):
        logger.error(f"转换后的文档不存在: {converted_path}")
        return jsonify({'success': False, 'msg': '转换后的文档不存在'}), 404

    logger.info(f"成功找到转换后的文档: {converted_path}")
    return send_file(converted_path, as_attachment=True, download_name='转换后_5_功能需求_Demo文档.docx')


@document_bp.route('/format-check', methods=['POST'])
def format_check():
    """
    检查上传文档与标准格式的匹配度
    """
    logger.info("开始文档格式检查")

    if 'document_file' not in request.files:
        logger.warning("未选择文件")
        return render_template('format_check.html',
                               result={"match_rate": 0, "errors": ["未选择文件"]})

    doc_file = request.files['document_file']
    if doc_file.filename == '':
        logger.warning("文件名为空")
        return render_template('format_check.html',
                               result={"match_rate": 0, "errors": ["文件名不能为空"]})

    # 保存临时文件
    temp_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'temp_check.docx')
    doc_file.save(temp_path)
    logger.info(f"临时文件已保存至: {temp_path}")

    # 标准模板路径
    standard_template_path = 'templates/功能需求_Demo模板.docx'
    # 执行格式匹配
    match_result = match_document_format(temp_path, standard_template_path)

    # 清理临时文件
    os.remove(temp_path)

    logger.info(f"格式检查完成，匹配度: {match_result['match_rate']:.2f}")
    return render_template('format_check.html', result=match_result)
@document_bp.route('/format-document', methods=['POST'])
def format_document():
    """处理文档格式化请求"""
    logger.info("开始处理文档格式化请求")

    try:
        # 检查模板是否存在
        template_path = 'templates/功能需求_Demo模板.docx'
        if not os.path.exists(template_path):
            logger.warning("Demo模板不存在")
            return render_template('document_formatter.html',
                                   format_result={'success': False, 'msg': '格式模板不存在'})

        # 处理上传文件
        if 'source_file' not in request.files:
            logger.warning("未选择待格式化文件")
            return render_template('document_formatter.html',
                                   format_result={'success': False, 'msg': '未选择待格式化文件'})

        source_file = request.files['source_file']
        if source_file.filename == '':
            logger.warning("待格式化文件名为空")
            return render_template('document_formatter.html',
                                   format_result={'success': False, 'msg': '文件名不能为空'})

        # 保存上传文件
        source_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'source_to_format.docx')
        source_file.save(source_path)
        logger.info(f"待格式化文件已保存至: {source_path}")

        # 分析模板格式
        logger.info("开始分析模板格式")
        template_format = analyze_template_format(template_path)

        # 应用格式到文档（使用增强版）
        output_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'formatted_document.docx')
        logger.info("开始应用格式到文档")

        success, error = enhanced_apply_format_to_document(source_path, template_format, output_path)

        if not success:
            logger.error(f"文档格式化失败: {error}")
            return render_template('document_formatter.html',
                                   format_result={'success': False, 'msg': f'文档格式化失败：{error}'})

        logger.info("文档格式化成功")
        return render_template('document_formatter.html',
                               format_result={'success': True, 'msg': '文档格式化成功',
                                              'download_url': '/download-formatted-doc'})

    except Exception as e:
        logger.error(f"文档格式化过程中发生未预期错误: {str(e)}", exc_info=True)
        return render_template('document_formatter.html',
                               format_result={'success': False, 'msg': f'系统错误：{str(e)}'})

@document_bp.route('/download-formatted-doc')
def download_formatted_doc():
    """下载格式化后的文档"""
    logger.info("开始处理格式化文档下载请求")
    formatted_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'formatted_document.docx')
    if not os.path.exists(formatted_path):
        logger.error(f"格式化后的文档不存在: {formatted_path}")
        return jsonify({'success': False, 'msg': '格式化后的文档不存在'}), 404

    logger.info(f"成功找到格式化后的文档: {formatted_path}")
    return send_file(formatted_path, as_attachment=True,
                     download_name='格式化后的文档.docx')


@document_bp.route('/document-formatter')
def document_formatter():
    """文档格式化页面"""
    logger.info("访问文档格式化页面")
    template_exists = os.path.exists('templates/功能需求_Demo模板.docx')
    logger.info(f"模板存在状态: {template_exists}")
    return render_template('document_formatter.html',
                           template_exists=template_exists,
                           format_result=None)
