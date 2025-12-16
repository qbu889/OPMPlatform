import os
import logging
import re

from flask import Blueprint, request, render_template, send_file, json
from werkzeug.utils import secure_filename
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from bs4 import BeautifulSoup

# 创建 Blueprint 实例
markdown_upload_bp = Blueprint('markdown', __name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'md'}

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 设置日志配置，默认INFO级别
logging.basicConfig(level=logging.INFO)

KEYWORDS_FILE = os.path.join('config', 'keywords.json')
TEMPLATE_DOCX = "templates/template.docx"

# 关键字段列表，包含常见变体，方便后续匹配
FIELD_NAMES = [
    "场景说明：",
    "功能描述：",
    "系统界面：",
    "输入：",
    "输出：",
    "处理过程：",
    "本事务功能涉及到的数据文件（即FTR/RET）：",
    "本事务功能涉及到的数据文件（即 FTR/RET）：",
    "本事务功能涉及到的数据文件（即FTR/RET）:",
    "本事务功能涉及到的数据文件（即 FTR/RET）:",
    "本事务功能涉及到的数据文件（即FTR/RET）",
    "本事务功能涉及到的数据文件（即 FTR/RET）",
]

def allowed_file(filename: str) -> bool:
    """判断文件后缀是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@markdown_upload_bp.route('/markdown-upload', methods=['GET', 'POST'])
def markdown_upload():
    """
    Markdown文件上传与转换接口
    支持GET显示上传页面，POST上传文件并转换为Word后返回下载
    """
    try:
        logging.info("请求 /markdown-upload 方法: %s", request.method)
        if request.method == 'POST':
            if 'markdown_file' not in request.files:
                logging.warning("POST请求缺少markdown_file")
                return render_template('markdown_upload.html', error='没有选择文件')

            file = request.files['markdown_file']
            if file.filename == '' or not allowed_file(file.filename):
                logging.warning("上传文件无效或为空，文件名: %s", file.filename)
                return render_template('markdown_upload.html', error='没有选择有效的 Markdown 文件')

            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)

            # 防止路径穿越攻击
            real_path = os.path.realpath(filepath)
            upload_real_path = os.path.realpath(UPLOAD_FOLDER)
            if not real_path.startswith(upload_real_path):
                logging.error("非法路径访问尝试，路径: %s", real_path)
                raise ValueError("非法路径访问尝试")

            file.save(real_path)
            logging.info("上传文件保存成功: %s", real_path)

            word_filename = clean_filename(file.filename).replace('.md', '.docx').replace('.MD', '.docx')
            word_filepath = os.path.join(UPLOAD_FOLDER, word_filename)

            logging.info("开始转换Markdown为Word: %s -> %s", real_path, word_filepath)
            convert_md_to_docx(real_path, word_filepath)
            logging.info("转换完成，准备发送文件: %s", word_filepath)

            response = send_file(word_filepath, as_attachment=True)

            @response.call_on_close
            def cleanup():
                try:
                    os.remove(real_path)
                    os.remove(word_filepath)
                    logging.info("临时文件删除成功：%s 和 %s", real_path, word_filepath)
                except Exception as e:
                    logging.warning(f"Cleanup failed: {e}")

            return response

        # GET请求显示上传页面
        return render_template('markdown_upload.html')

    except Exception as e:
        logging.error(f"文件处理异常: {e}", exc_info=True)
        return render_template('markdown_upload.html', error="处理过程中发生错误，请稍后再试")


def set_font(run, name='宋体', size=None, bold=False):
    """
    设置Run对象字体样式
    """
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色字体


def clean_filename(filename: str) -> str:
    """
    清理文件名，替换非法字符，避免文件系统错误
    """
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename.strip(' ._')  # 去除首尾空格和点


def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    md_content = remove_scenario_description_lines(md_content)
    md_content = remove_fpa_report_lines(md_content)  # 新增：移除包含“FPA功能点分析报告”的行
    md_content = update_file_statistics(md_content)
    md_content = process_function_description(md_content)
    md_content = re.sub(r'^(#{7,})', '######', md_content, flags=re.MULTILINE)

    html = markdown.markdown(md_content, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')

    doc = Document()
    # 添加固定功能需求描述，避免重复添加
    # add_fixed_function_requirements(doc)
    # 解析HTML元素，生成对应Word内容
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = doc.add_heading(level=min(level, 9))
            text = element.get_text().strip()
            if text:
                # 去除标题序号
                text = re.sub(r'^\d+(\.\d+)*\.\s*', '', text)
                run = heading.add_run(text)
                set_font(run, size=get_heading_size(level), bold=True)
            continue

        if element.name == 'p':
            text = element.get_text().strip()
            if not text:
                continue  # 跳过空段落
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, element)
            continue

        if element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                li_text = li.get_text().strip()
                if not li_text:
                    continue
                paragraph = doc.add_paragraph(style='List Bullet')
                add_formatted_text(paragraph, li)



    # 保存Word文档
    doc.save(docx_path)
    logging.info("Word文档保存成功: %s", docx_path)


def remove_scenario_description_lines(content: str) -> str:
    """
    删除所有包含“场景说明”的行
    """
    lines = content.split('\n')
    filtered_lines = [line for line in lines if '场景说明' not in line]
    return '\n'.join(filtered_lines)


def remove_fpa_report_lines(content: str) -> str:
    """
    删除所有包含“FPA功能点分析报告”的行
    """
    lines = content.split('\n')
    filtered_lines = [line for line in lines if 'FPA功能点分析报告' not in line]
    return '\n'.join(filtered_lines)


def remove_numbering_from_process_line(line: str) -> str:
    """
    去除“处理过程”字段中数字序号，支持Markdown加粗格式
    """
    prefix_match = re.match(r'^\**处理过程[:：]\**\s*(.*)$', line)
    if not prefix_match:
        return line

    content = prefix_match.group(1)
    parts = re.split(r'[；;]', content)
    cleaned_parts = [re.sub(r'^\s*\d+\.\s*', '', part).strip() for part in parts if part.strip()]
    new_content = '；'.join(cleaned_parts)

    return f'**处理过程：**{new_content}'


def process_function_description(content: str) -> str:
    """
    处理功能描述字段：
    - 将“处理xxx相关业务”修改为“进行xxx”
    - 根据关键词修改输入字段描述
    - 去除处理过程数字序号
    """
    logging.info("开始处理功能描述")

    keywords_config = load_keywords_config()
    person_keywords = keywords_config.get('person_keywords', [])
    system_keywords = keywords_config.get('system_keywords', [])

    # pattern_title = r'^(#{6,7})\s*\*\*(.*?)\*\*$'
    pattern_title = r'^(#{5,7})\s*(?:\*\*)?(.*?)(?:\*\*)?\s*$'
    lines = content.split('\n')

    func_desc_updated_count = 0
    input_updated_count = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        # 处理“处理过程”去序号
        if re.match(r'^\**处理过程[:：]\**', line):
            lines[i] = remove_numbering_from_process_line(line)
            line = lines[i]

        match = re.match(pattern_title, line)
        if match:
            full_title = match.group(2).strip()
            number_prefix_match = re.match(r'^([\d\.]+)\s+(.*)$', full_title)
            func_name = number_prefix_match.group(2).strip() if number_prefix_match else full_title

            contains_keyword = any(keyword in func_name for keyword in person_keywords)
            contains_keyword2 = any(keyword in func_name for keyword in system_keywords)

            j = i + 1
            desc_processed = False
            input_processed = False

            while j < len(lines):
                desc_line = lines[j]
                if re.match(r'^#{1,7}\s*', desc_line):
                    break

                if not desc_processed:
                    desc_pattern = r'^(\**功能描述[:：]\**)\s*(.*)$'
                    desc_match = re.match(desc_pattern, desc_line)
                    if desc_match:
                        prefix = desc_match.group(1)
                        new_desc = '进行' + func_name
                        lines[j] = f"{prefix} {new_desc}"
                        desc_processed = True
                        func_desc_updated_count += 1

                if (contains_keyword or contains_keyword2) and not input_processed:
                    input_pattern = r'^(\**输入[:：]\**)\s*(.*)$'
                    input_match = re.match(input_pattern, desc_line)
                    if input_match:
                        prefix = input_match.group(1)
                        new_input = ('用户触发' if contains_keyword else '系统判定') + func_name
                        lines[j] = f"{prefix} {new_input}"
                        input_processed = True
                        input_updated_count += 1

                j += 1
        i += 1

    logging.info("功能描述更新数量: %d, 输入字段更新数量: %d", func_desc_updated_count, input_updated_count)
    logging.info("功能描述处理完成")
    return '\n'.join(lines)


from docx.shared import Pt

def add_fixed_function_requirements(doc: Document):
    existing_content = "\n".join(p.text for p in doc.paragraphs)
    if "按照FPA中列出的本需求需改造的功能逐级" in existing_content or "功能需求" in existing_content:
        return

    heading = doc.add_heading('功能需求', level=1)
    set_font(heading.runs[0], size=16, bold=True)

    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run('按照')
    run1.bold = True
    set_font(run1, size=12, bold=True)

    run2 = paragraph.add_run('FPA')
    run2.bold = True
    set_font(run2, size=12, bold=True)

    run3 = paragraph.add_run('中列出的本需求需改造的功能逐级（即按一级分类、二级分类、三级分类、功能点名称、功能点计数项结构）描述功能需求。')
    set_font(run3, size=12, bold=False)


def add_formatted_text(paragraph, element):
    """
    递归处理HTML元素，保留加粗和换行，去除斜体格式
    """
    if element.name is None:
        text = str(element)
        if text.strip() or text.isspace():
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if line.strip():
                    run = paragraph.add_run(line)
                    set_font(run, size=12)
                if i < len(lines) - 1:
                    paragraph.add_run('\n')
        return

    for child in element.children:
        if hasattr(child, 'name'):
            if child.name in ['strong', 'b']:
                text = child.get_text()
                lines = text.split('\n')
                for i, line in enumerate(lines):
                    if line.strip():
                        run = paragraph.add_run(line)
                        run.bold = True
                        set_font(run, size=12, bold=True)
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')
            elif child.name in ['em', 'i']:
                text = child.get_text()
                lines = text.split('\n')
                for i, line in enumerate(lines):
                    if line.strip():
                        run = paragraph.add_run(line)
                        set_font(run, size=12)  # 不应用斜体
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')
            elif child.name == 'br':
                paragraph.add_run('\n')
            else:
                add_formatted_text(paragraph, child)
        else:
            text = str(child)
            if text.strip() or text.isspace():
                lines = text.split('\n')
                for i, line in enumerate(lines):
                    if line.strip():
                        run = paragraph.add_run(line)
                        set_font(run, size=12)
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')


def count_files_by_separators(text: str) -> int:
    """
    根据逗号、分号、空格、换行等分隔符统计文件数量，'无'或空返回0
    """
    if not text or text == '无':
        return 0
    parts = re.split(r'[，,；;\s]+', text.strip())
    parts = [p for p in parts if p]
    return len(parts)


def update_file_statistics(content: str) -> str:
    """
    更新本事务功能预计涉及的数据文件统计内容
    格式化数字空格，计算实际文件数量并替换文本
    """
    cleaned_content = re.sub(r'(\d+)\s+个', r'\1个', content)
    cleaned_content = re.sub(r'新增\s*/\s*变更', '新增/变更', cleaned_content)
    cleaned_content = re.sub(r'[ \t]*：[ \t]*', '：', cleaned_content)

    pattern = r'本事务功能预计涉及到\s*(\d+)\s*个内部逻辑文件\s*，?\s*(\d+)\s*个外部逻辑文件'
    match = re.search(pattern, cleaned_content)

    if not match:
        return cleaned_content

    # 内部逻辑文件统计
    internal_new_added_pattern = r'本期新增/变更的内部逻辑文件[：:]([^\n]*?)(?=\n|$)'
    internal_existing_pattern = r'本期涉及原有但没修改的内部逻辑文件[：:]([^\n]*?)(?=\n|$)'

    internal_new_added_match = re.search(internal_new_added_pattern, cleaned_content)
    internal_existing_match = re.search(internal_existing_pattern, cleaned_content)

    internal_new_count = count_files_by_separators(internal_new_added_match.group(1).strip()) if internal_new_added_match else 0
    internal_existing_count = count_files_by_separators(internal_existing_match.group(1).strip()) if internal_existing_match else 0
    actual_internal_count = internal_new_count + internal_existing_count

    # 外部逻辑文件统计
    external_new_added_pattern = r'本期新增/变更的外部逻辑文件[：:]([^\n]*?)(?=\n|$)'
    external_existing_pattern = r'本期涉及原有但没修改的外部逻辑文件[：:]([^\n]*?)(?=\n|$)'

    external_new_added_match = re.search(external_new_added_pattern, cleaned_content)
    external_existing_match = re.search(external_existing_pattern, cleaned_content)

    external_new_count = count_files_by_separators(external_new_added_match.group(1).strip()) if external_new_added_match else 0
    external_existing_count = count_files_by_separators(external_existing_match.group(1).strip()) if external_existing_match else 0
    actual_external_count = external_new_count + external_existing_count

    updated_content = re.sub(
        pattern,
        f'本事务功能预计涉及到{actual_internal_count}个内部逻辑文件，{actual_external_count}个外部逻辑文件',
        cleaned_content,
        count=1
    )
    return updated_content


def init_keywords_config():
    """
    初始化关键词配置文件，首次运行时创建默认配置
    """
    if not os.path.exists('config'):
        os.makedirs('config')
        logging.info("创建config目录")

    if not os.path.exists(KEYWORDS_FILE):
        default_config = {
            "person_keywords": [
                "增加", "删除", "修改", "查询", "呈现", "列表", "渲染", "导出", "页面",
                "自定义", "点击", "看板", "录入", "编辑", "查看", "搜索", "筛选", "排序",
                "下载", "上传", "选择", "勾选", "取消", "确认", "提交", "保存", "刷新",
                "切换", "拖拽", "复制", "粘贴", "撤销", "重做", "下钻", "钻取", "展开",
                "收缩", "过滤", "钻探", "跳转", "关联", "流转", "设置", "控制"
            ],
            "system_keywords": [
                "判断", "自动", "定时", "轮询", "监听", "检测", "校验", "验证", "同步",
                "异步", "回调", "触发", "推送", "通知", "告警", "计算", "统计", "分析",
                "识别", "匹配", "对比", "评估", "鉴权", "认证", "授权", "加密", "解密",
                "签名", "验签", "审计", "监控", "日志", "备份", "恢复", "缓存", "对接",
                "调用", "判定", "记录"
            ]
        }
        with open(KEYWORDS_FILE, 'w', encoding='utf-8') as f:
            json.dump(default_config, f, ensure_ascii=False, indent=2)
        logging.info("默认关键词配置文件已创建")


def load_keywords_config():
    """
    加载关键词配置，若无配置则初始化
    """
    if not os.path.exists(KEYWORDS_FILE):
        logging.info("关键词配置文件不存在，初始化默认配置")
        init_keywords_config()

    with open(KEYWORDS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_keywords_config(config):
    """
    保存关键词配置到文件
    """
    with open(KEYWORDS_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


@markdown_upload_bp.route('/api/keywords', methods=['GET'])
def get_keywords_api():
    """
    获取关键词配置API
    """
    try:
        config = load_keywords_config()
        return config
    except Exception as e:
        logging.error("获取关键词配置失败: %s", e)
        return {'status': 'error', 'message': str(e)}, 500


@markdown_upload_bp.route('/api/keywords', methods=['POST'])
def update_keywords_api():
    """
    更新关键词配置API
    支持添加和编辑关键词
    """
    try:
        data = request.json
        keyword_type = data.get('type')
        keyword = data.get('keyword')
        original = data.get('original')

        type_mapping = {'person': 'person_keywords', 'system': 'system_keywords'}
        if keyword_type in type_mapping:
            keyword_type = type_mapping[keyword_type]

        if not keyword_type or not keyword:
            return {'status': 'error', 'message': '缺少必要参数'}, 400

        config = load_keywords_config()

        if original:
            # 编辑关键词，先删除原有，再添加新词
            if original in config.get(keyword_type, []):
                config[keyword_type].remove(original)
            config[keyword_type].append(keyword)
        else:
            # 新增关键词，避免重复添加
            if keyword not in config.get(keyword_type, []):
                config[keyword_type].append(keyword)

        save_keywords_config(config)
        return {'status': 'success'}
    except Exception as e:
        logging.error("更新关键词配置失败: %s", e)
        return {'status': 'error', 'message': str(e)}, 500


@markdown_upload_bp.route('/api/keywords', methods=['DELETE'])
def delete_keywords_api():
    """
    删除关键词配置API
    """
    try:
        data = request.json
        keyword_type = data.get('type')
        keyword = data.get('keyword')

        type_mapping = {'person': 'person_keywords', 'system': 'system_keywords'}
        if keyword_type in type_mapping:
            keyword_type = type_mapping[keyword_type]

        if not keyword_type or not keyword:
            return {'status': 'error', 'message': '缺少必要参数'}, 400

        config = load_keywords_config()

        if keyword_type not in config:
            return {'status': 'error', 'message': '无效的关键词类型'}, 400

        if keyword in config[keyword_type]:
            config[keyword_type].remove(keyword)
            save_keywords_config(config)

        return {'status': 'success'}
    except Exception as e:
        logging.error("删除关键词配置失败: %s", e)
        return {'status': 'error', 'message': str(e)}, 500


def get_heading_size(level: int) -> int:
    """
    根据标题级别返回对应字体大小，默认12pt
    """
    sizes = {
        1: 16,
        2: 15,
        3: 14,
        4: 12,
        5: 12,
        6: 12,
        7: 12  # 支持7级标题
    }
    return sizes.get(level, 12)
