import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from markdown import markdown
from bs4 import BeautifulSoup  # 需要：pip install beautifulsoup4

def load_markdown_as_blocks(md_path: str):
    """将 Markdown 转成结构化块（段落/标题/列表等）"""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 转成 HTML
    html = markdown(md_text, extensions=['fenced_code', 'tables'])
    soup = BeautifulSoup(html, 'html.parser')

    blocks = []
    for elem in soup.children:
        if elem.name is None:
            continue
        if re.match(r'h[1-6]', elem.name):
            level = int(elem.name[1])
            blocks.append({
                'type': 'heading',
                'level': level,
                'text': elem.get_text(strip=True)
            })
        elif elem.name == 'p':
            blocks.append({
                'type': 'paragraph',
                'text': elem.get_text(strip=True)
            })
        elif elem.name in ('ul', 'ol'):
            list_type = 'bullet' if elem.name == 'ul' else 'number'
            for li in elem.find_all('li', recursive=False):
                blocks.append({
                    'type': 'list_item',
                    'list_type': list_type,
                    'text': li.get_text(strip=True)
                })
        elif elem.name == 'pre':
            code = elem.get_text()
            blocks.append({
                'type': 'code',
                'text': code
            })
        elif elem.name == 'table':
            rows = []
            for tr in elem.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['td','th'])]
                rows.append(cells)
            blocks.append({
                'type': 'table',
                'rows': rows
            })
    return blocks

def fill_template_with_blocks(template_path: str, md_path: str, output_path: str):
    doc = Document(template_path)
    blocks = load_markdown_as_blocks(md_path)

    # 先尝试找占位符
    content_paragraph = None
    for para in doc.paragraphs:
        if '{{CONTENT}}' in para.text:
            content_paragraph = para
            break

    if content_paragraph is not None:
        # 找到了占位符，就按原来的逻辑：
        parent = content_paragraph._p.getparent()
        idx = parent.index(content_paragraph._p)
        parent.remove(content_paragraph._p)
    else:
        # 没找到，占位符就当不存在：把内容插在文档最后
        last_para = doc.paragraphs[-1]
        parent = last_para._p.getparent()
        idx = parent.index(last_para._p) + 1

    def add_paragraph_with_style(text, style_name):
        p = doc.add_paragraph(text)
        p.style = style_name
        parent.remove(p._p)
        parent.insert(idx, p._p)

    for block in blocks:
        if block['type'] == 'heading':
            level = block['level']
            if level == 1:
                style_name = 'Heading 1'
            elif level == 2:
                style_name = 'Heading 2'
            else:
                style_name = 'Heading 3'
            add_paragraph_with_style(block['text'], style_name)
            idx += 1

        elif block['type'] == 'paragraph':
            add_paragraph_with_style(block['text'], 'BodyText')
            idx += 1

        elif block['type'] == 'list_item':
            if block['list_type'] == 'bullet':
                style_name = 'Bullet List'  # 替换为你模板中实际存在的样式名称
            else:
                style_name = 'List Number'  # 替换为你模板中实际存在的样式名称
            add_paragraph_with_style(block['text'], style_name)
            idx += 1

        elif block['type'] == 'code':
            add_paragraph_with_style(block['text'], 'Code')
            idx += 1

        elif block['type'] == 'table':
            table = doc.add_table(rows=0, cols=len(block['rows'][0]))
            table.style = 'Table Grid'
            parent.remove(table._element)
            parent.insert(idx, table._element)
            for row in block['rows']:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row):
                    row_cells[i].text = cell_text
            idx += 1

    doc.save(output_path)

if __name__ == '__main__':
    fill_template_with_blocks(
        template_path='/Users/linziwang/PycharmProjects/wordToWord/templates/template.docx',
        md_path='/Users/linziwang/PycharmProjects/wordToWord/static/12091510集中故障管理系统-监控综合应用-关于集团事件工单省部接口数据上报保障的开发需求-15400_20250318.md',
        output_path='output.docx'
    )
