"""
Excel 转 COSMIC Word 文档路由
基于 convert_script 项目改造
"""
import os
import time
from pathlib import Path
from flask import Blueprint, request, jsonify, send_file
from werkzeug.utils import secure_filename
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging

logger = logging.getLogger(__name__)

cosmic_bp = Blueprint('cosmic', __name__, url_prefix='/api/cosmic')

# 配置
UPLOAD_FOLDER = Path('uploads/cosmic')
OUTPUT_FOLDER = Path('downloads/cosmic')
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}

# 确保目录存在
UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)


def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def set_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置字体格式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def get_font_size_for_level(level):
    """根据标题级别返回字号"""
    if level == 3: return 16
    if level == 4: return 14
    if level == 5: return 12
    if level == 6: return 10.5
    return 10.5


def split_subprocess_description(text):
    """拆分子过程描述字段"""
    if not text or pd.isna(text):
        return []

    text = str(text).strip()
    if not text:
        return []

    prefixes = ['输入-', '查询-', '呈现-', '校验-', '输出-']
    result = []
    current_segment = []

    i = 0
    while i < len(text):
        is_prefix = False
        for prefix in prefixes:
            if text[i:i+len(prefix)] == prefix:
                if current_segment:
                    result.append(''.join(current_segment).strip())
                    current_segment = []
                current_segment.append(prefix)
                i += len(prefix)
                is_prefix = True
                break

        if not is_prefix:
            current_segment.append(text[i])
            i += 1

    if current_segment:
        result.append(''.join(current_segment).strip())

    result = [seg.rstrip('；;').strip() for seg in result if seg]

    if not result and text:
        return [text]

    return result


def read_excel_robust(excel_path):
    """健壮地读取 Excel 文件"""
    try:
        xl = pd.ExcelFile(excel_path)
    except Exception as e:
        logger.error(f"无法打开 Excel 文件：{e}")
        return None

    # 查找包含数据的 Sheet
    target_sheet = None
    for sheet in xl.sheet_names:
        if '拆分表' in sheet or '功能点' in sheet:
            target_sheet = sheet
            break

    if not target_sheet:
        target_sheet = xl.sheet_names[0]
        logger.warning(f"未找到特定 Sheet，使用：{target_sheet}")
    else:
        logger.info(f"使用 Sheet: {target_sheet}")

    # Excel 结构分析 (0-indexed):
    # 行 0: 通用软件评估模型
    # 行 1: 度量策略阶段 | 映射阶段 | 度量阶段
    # 行 2: 客户需求 | 功能用户 | 触发事件 | 功能过程 | 子过程描述 | ...
    # 行 3: NaN | 一级模块 | 二级模块 | 三级模块 (这些列标签)
    # 行 4+: 数据行 - 列 1=一级模块值，列 2=二级模块值，列 3=三级模块值，列 5=功能过程，列 7=子过程描述

    df_raw = pd.read_excel(excel_path, sheet_name=target_sheet, header=None)

    # 查找"一级模块" | "二级模块" | "三级模块" 所在的行 (第 3 行)
    header_row_idx = 3
    for idx, row in df_raw.iloc[:5].iterrows():
        row_values = [str(v) if pd.notna(v) else '' for v in row]
        row_text = ' '.join(row_values)
        if '一级模块' in row_text and '二级模块' in row_text and '三级模块' in row_text:
            header_row_idx = int(idx)
            break

    logger.info(f"定位到表头在第 {header_row_idx} 行")

    # 数据从第 header_row_idx + 1 行开始
    df_data = df_raw.iloc[header_row_idx + 1:].reset_index(drop=True)

    # 数据行的列索引：
    # 列 1: 一级模块值 (如"流程管理")
    # 列 2: 二级模块值 (如"数据接入流程")
    # 列 3: 三级模块值 (如"数据接入流程")
    # 列 6: 功能过程 (如"创建数据接入流程工单") - 注意：列 5 是触发事件
    # 列 7: 子过程描述 (如"输入-用户...")

    module_l1_col = 1
    module_l2_col = 2
    module_l3_col = 3
    func_col = 6       # 修正：功能过程在列 6，不是列 5
    subproc_col = 7

    # 收集数据 - 每行的模块值直接读取
    result = []
    for idx, row in df_data.iterrows():
        values = row.tolist()

        # 从数据行直接读取模块值
        module_l1 = str(values[module_l1_col]).strip() if module_l1_col < len(values) and pd.notna(values[module_l1_col]) else ''
        module_l2 = str(values[module_l2_col]).strip() if module_l2_col < len(values) and pd.notna(values[module_l2_col]) else ''
        module_l3 = str(values[module_l3_col]).strip() if module_l3_col < len(values) and pd.notna(values[module_l3_col]) else ''

        # 获取功能过程和子过程描述
        function_name = str(values[func_col]).strip() if func_col < len(values) and pd.notna(values[func_col]) else ''
        subprocess_desc = str(values[subproc_col]).strip() if subproc_col < len(values) and pd.notna(values[subproc_col]) else ''

        # 不再在这里过滤，交给 excel_to_word_conversion 处理（需要向下填充）
        result.append({
            '一级模块': module_l1,
            '二级模块': module_l2,
            '三级模块': module_l3,
            '功能过程': function_name,
            '子过程描述': subprocess_desc
        })

    return pd.DataFrame(result)

def excel_to_word_conversion(excel_path, word_path):
    """执行 Excel 到 Word 的转换（COSMIC 标准格式）"""
    logger.info(f"正在处理：{excel_path.name}")

    df = read_excel_robust(excel_path)
    if df is None:
        raise Exception("无法读取 Excel 文件")

    # 创建 Word 文档
    doc = Document()

    # 设置默认字体为宋体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(10.5)  # 五号字

    # 按模块分组收集数据
    # 结构：{三级模块：{功能过程：[子过程描述]}}
    module_data = {}
    module_l1_value = None  # 一级模块值
    module_l2_value = None  # 二级模块值
    module_l3_value = None  # 三级模块值（用于向下填充）

    # Excel 数据结构说明 (基于 header=None 读取):
    # 行 0-2: 表头信息
    # 行 3: 包含"一级模块"、"二级模块"、"三级模块"等标签
    # 行 4+: 实际数据
    #   列 0: 客户需求描述
    #   列 1: 一级模块 (如"流程管理")
    #   列 2: 二级模块 (如"数据接入流程")
    #   列 3: 三级模块 (如"数据接入流程")
    #   列 4: 功能用户
    #   列 5: 触发事件
    #   列 6: 功能过程 (如"创建数据接入流程工单") ← 关键!
    #   列 7: 子过程描述 (如"输入-用户...")
    #   列 8+: 其他列
    
    # 直接使用固定列索引，不通过列名查找
    module_l1_col = 1   # 一级模块在列 1
    module_l2_col = 2   # 二级模块在列 2
    module_l3_col = 3   # 三级模块在列 3
    function_col = 6    # 功能过程在列 6 (不是列 5!)
    subprocess_col = 7  # 子过程描述在列 7

    logger.info(f"使用列索引：一级模块={module_l1_col}, 二级模块={module_l2_col}, 三级模块={module_l3_col}, 功能过程={function_col}, 子过程描述={subprocess_col}")

    # 调试：统计处理的数据行
    total_rows = len(df)
    processed_rows = 0
    skipped_rows = 0
    function_name_value = None  # 功能过程值（用于向下填充）
    logger.debug(f"开始处理 {total_rows} 行数据")

    for idx, row in df.iterrows():
        # 提取模块信息（通过列名）
        module_l1 = str(row.get('一级模块', '')) if pd.notna(row.get('一级模块', '')) else ''
        module_l2 = str(row.get('二级模块', '')) if pd.notna(row.get('二级模块', '')) else ''
        module_l3 = str(row.get('三级模块', '')) if pd.notna(row.get('三级模块', '')) else ''
        function_name = str(row.get('功能过程', '')) if pd.notna(row.get('功能过程', '')) else ''
        subprocess_desc = str(row.get('子过程描述', '')) if pd.notna(row.get('子过程描述', '')) else ''

        # 记录一级模块和二级模块（取第一个非空值）
        if module_l1 and module_l1 != 'nan' and module_l1.strip():
            module_l1_value = module_l1
            logger.debug(f"行 {idx}: 检测到一级模块: {module_l1}")
        if module_l2 and module_l2 != 'nan' and module_l2.strip():
            module_l2_value = module_l2
            logger.debug(f"行 {idx}: 检测到二级模块: {module_l2}")
        
        # 三级模块向下填充：如果当前行为空，使用上一个非空值
        if module_l3 and module_l3 != 'nan' and module_l3.strip():
            module_l3_value = module_l3
            logger.debug(f"行 {idx}: 检测到三级模块: {module_l3}")
        else:
            module_l3 = module_l3_value if module_l3_value else ''
        
        # 功能过程向下填充：如果当前行为空，使用上一个非空值
        if function_name and function_name != 'nan' and function_name.strip():
            function_name_value = function_name
            logger.debug(f"行 {idx}: 检测到功能过程: {function_name}")
        else:
            function_name = function_name_value if function_name_value else ''

        # 跳过空行（功能过程和子过程描述都为空）
        if (not function_name or function_name == 'nan') and (not subprocess_desc or subprocess_desc == 'nan'):
            skipped_rows += 1
            continue
        
        # 如果功能过程为空但子过程描述不为空，说明这行只有子过程描述，需要继承功能过程
        if not function_name or function_name == 'nan':
            # 这种情况不应该发生，因为上面已经做了向下填充
            # 但如果真的发生了，说明前面没有功能过程定义，跳过
            skipped_rows += 1
            logger.debug(f"行 {idx}: 跳过 - 功能过程为空但有子过程描述")
            continue

        processed_rows += 1
        logger.debug(f"行 {idx}: 功能过程={function_name}, 三级模块={module_l3}, 子过程={subprocess_desc[:50] if subprocess_desc else '无'}")

        # 按三级模块分组
        if module_l3 not in module_data:
            module_data[module_l3] = {}
            logger.debug(f"创建新的三级模块分组: {module_l3}")

        if function_name not in module_data[module_l3]:
            module_data[module_l3][function_name] = []
            logger.debug(f"添加新功能点: {function_name} 到模块 {module_l3}")

        # 收集子过程描述
        if subprocess_desc and subprocess_desc != 'nan' and subprocess_desc.strip():
            sub_items = split_subprocess_description(subprocess_desc)
            module_data[module_l3][function_name].extend(sub_items)
            logger.debug(f"添加 {len(sub_items)} 个子过程描述到 {function_name}")

    # 按一级模块 -> 三级模块分组（嵌套字典）
    # 结构: {一级模块: {三级模块: {功能过程: [子过程]}}}
    hierarchical_data = {}
    
    logger.info(f"数据处理完成: 总行数={total_rows}, 处理={processed_rows}, 跳过={skipped_rows}")
    logger.info(f"共识别到 {len(module_data)} 个三级模块")
    
    # 重新组织 module_data 为层级结构
    # 需要重新遍历，使用向下填充后的值
    module_l1_value = None
    module_l3_value = None
    function_name_value = None
    
    for idx, row in df.iterrows():
        module_l1 = str(row.get('一级模块', '')) if pd.notna(row.get('一级模块', '')) else ''
        module_l3 = str(row.get('三级模块', '')) if pd.notna(row.get('三级模块', '')) else ''
        function_name = str(row.get('功能过程', '')) if pd.notna(row.get('功能过程', '')) else ''
        subprocess_desc = str(row.get('子过程描述', '')) if pd.notna(row.get('子过程描述', '')) else ''
        
        # 向下填充
        if module_l1 and module_l1 != 'nan' and module_l1.strip():
            module_l1_value = module_l1
        if module_l3 and module_l3 != 'nan' and module_l3.strip():
            module_l3_value = module_l3
        if function_name and function_name != 'nan' and function_name.strip():
            function_name_value = function_name
        else:
            function_name = function_name_value if function_name_value else ''
        
        # 使用填充后的值
        if not module_l1 or module_l1 == 'nan':
            module_l1 = module_l1_value if module_l1_value else '未分类'
        if not module_l3 or module_l3 == 'nan':
            module_l3 = module_l3_value if module_l3_value else '未分类'
        
        # 跳过空行
        if not function_name or function_name == 'nan':
            continue
        
        # 构建层级结构
        if module_l1 not in hierarchical_data:
            hierarchical_data[module_l1] = {}
        
        if module_l3 not in hierarchical_data[module_l1]:
            hierarchical_data[module_l1][module_l3] = {}
        
        if function_name not in hierarchical_data[module_l1][module_l3]:
            hierarchical_data[module_l1][module_l3][function_name] = []
        
        # 添加子过程
        if subprocess_desc and subprocess_desc != 'nan' and subprocess_desc.strip():
            sub_items = split_subprocess_description(subprocess_desc)
            hierarchical_data[module_l1][module_l3][function_name].extend(sub_items)
    
    logger.info(f"层级数据统计:")
    for l1, l3_dict in hierarchical_data.items():
        logger.info(f"  一级模块 '{l1}' 包含 {len(l3_dict)} 个三级模块: {list(l3_dict.keys())}")
        for l3, funcs in l3_dict.items():
            logger.info(f"    三级模块 '{l3}' 包含 {len(funcs)} 个功能点: {list(funcs.keys())[:3]}...")

    # 生成文档
    # 标题 1：功能需求（固定）
    p = doc.add_heading('功能需求', level=1)
    for run in p.runs:
        set_font(run, font_name='宋体', font_size=22, bold=True)

    # 遍历每个一级模块 -> 三级模块
    for module_l1, l3_dict in hierarchical_data.items():
        # 标题 2：一级模块（流程管理、国产化改造等）
        p = doc.add_heading(module_l1, level=2)
        for run in p.runs:
            set_font(run, font_name='宋体', font_size=18, bold=True)
        
        # 遍历该一级模块下的所有三级模块
        for module_l3, functions in l3_dict.items():
            # 标题 3：二级模块/功能分类（如"数据接入流程"）
            p = doc.add_heading(module_l3, level=3)
            for run in p.runs:
                set_font(run, font_name='宋体', font_size=16, bold=True)
                    
            # 标题 4：三级模块 = H3 + "（新增）"后缀（如"数据接入流程（新增）"）
            h4_title = f"{module_l3}（新增）"
            p = doc.add_heading(h4_title, level=4)
            for run in p.runs:
                set_font(run, font_name='宋体', font_size=14, bold=True)
        
            # 标题 5：关键时序图/业务逻辑图（固定）
            p = doc.add_heading('关键时序图/业务逻辑图', level=5)
            for run in p.runs:
                set_font(run, font_name='宋体', font_size=12, bold=True)
        
            # 正文：无
            p = doc.add_paragraph('无。')
            for run in p.runs:
                set_font(run, font_name='宋体', font_size=10.5, bold=False)
            # 通过 XML 直接设置左缩进 2 字符（800 twips = 40 磅）
            from docx.oxml import OxmlElement
            pPr = p._element.get_or_add_pPr()
            old_ind = pPr.find(qn('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '800')
            pPr.insert(0, ind)
        
            # 标题 5：功能描述（固定）
            p = doc.add_heading('功能描述', level=5)
            for run in p.runs:
                set_font(run, font_name='宋体', font_size=12, bold=True)
        
            # 整体功能列表（该三级模块下的所有功能过程）
            all_functions = list(functions.keys())
            p = doc.add_paragraph('整体功能列表包含如下：' + '、'.join(all_functions) + '。')
            for run in p.runs:
                set_font(run, font_name='宋体', font_size=10.5, bold=False)
            # 通过 XML 直接设置首行缩进 2 字符（800 twips = 40 磅）
            from docx.oxml import OxmlElement
            pPr = p._element.get_or_add_pPr()
            # 清除可能存在的 ind 元素
            old_ind = pPr.find(qn('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            # 创建新的 ind 元素，只设置 firstLine
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '800')
            pPr.insert(0, ind)
        
            # 详细功能点描述（带序号）- 按照 cosmic.md 的要求格式
            for i, (func_name, subprocesses) in enumerate(functions.items(), 1):
                # 功能点标题（编号 + 功能过程名称）- 不需要加粗
                p = doc.add_paragraph()
                run = p.add_run(f"{i}.{func_name}")
                set_font(run, font_name='宋体', font_size=10.5, bold=False)
                # 通过 XML 直接设置首行缩进 2 字符（800 twips = 40 磅）
                from docx.oxml import OxmlElement
                pPr = p._element.get_or_add_pPr()
                # 清除可能存在的 ind 元素
                old_ind = pPr.find(qn('w:ind'))
                if old_ind is not None:
                    pPr.remove(old_ind)
                # 创建新的 ind 元素，只设置 firstLine
                ind = OxmlElement('w:ind')
                ind.set(qn('w:firstLine'), '800')
                pPr.insert(0, ind)
        
                # 子过程描述（首行缩进）
                for subprocess in subprocesses:
                    p = doc.add_paragraph(subprocess)
                    for run in p.runs:
                        set_font(run, font_name='宋体', font_size=10.5, bold=False)
                    # 通过 XML 直接设置首行缩进 2 字符（800 twips = 40 磅）
                    from docx.oxml import OxmlElement
                    pPr = p._element.get_or_add_pPr()
                    # 清除可能存在的 ind 元素
                    old_ind = pPr.find(qn('w:ind'))
                    if old_ind is not None:
                        pPr.remove(old_ind)
                    # 创建新的 ind 元素，只设置 firstLine
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:firstLine'), '800')
                    pPr.insert(0, ind)
            # 添加空行分隔
            doc.add_paragraph()

    # 保存文档
    doc.save(word_path)
    logger.info(f"Word 文档已保存：{word_path}")

    # 统计功能点数量
    total_functions = sum(
        len(funcs) 
        for l1_dict in hierarchical_data.values() 
        for funcs in l1_dict.values()
    )
    return total_functions


@cosmic_bp.route('/upload', methods=['POST'])
def upload_excel():
    """上传 Excel 文件"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '文件名为空'}), 400

        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': '不支持的文件格式，请上传 .xlsx 或 .xls 文件'}), 400

        # 生成唯一文件名
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        unique_filename = f"{Path(original_filename).stem}_{timestamp}{Path(original_filename).suffix}"

        # 保存文件
        file_path = UPLOAD_FOLDER / unique_filename
        file.save(file_path)

        logger.info(f"文件上传成功：{unique_filename}")

        return jsonify({
            'success': True,
            'message': '文件上传成功',
            'filename': unique_filename,
            'original_name': original_filename
        })

    except Exception as e:
        logger.error(f"文件上传失败：{e}", exc_info=True)
        return jsonify({'success': False, 'message': f'上传失败：{str(e)}'}), 500


@cosmic_bp.route('/convert', methods=['POST'])
def convert_excel():
    """将 Excel 转换为 Word"""
    try:
        data = request.json
        filename = data.get('filename')

        if not filename:
            return jsonify({'success': False, 'message': '缺少文件名参数'}), 400

        excel_path = UPLOAD_FOLDER / filename
        if not excel_path.exists():
            return jsonify({'success': False, 'message': '文件不存在'}), 404

        # 读取数据用于统计
        df = read_excel_robust(excel_path)
        if df is None:
            return jsonify({'success': False, 'message': '无法读取 Excel 文件'}), 500

        # 统计模块数据
        module_l1_set = set()
        module_l2_set = set()
        module_l3_set = set()
        function_set = set()
        subprocess_count = 0
        
        module_l1_value = None
        module_l2_value = None
        module_l3_value = None
        function_name_value = None

        for idx, row in df.iterrows():
            module_l1 = str(row.get('一级模块', '')) if pd.notna(row.get('一级模块', '')) else ''
            module_l2 = str(row.get('二级模块', '')) if pd.notna(row.get('二级模块', '')) else ''
            module_l3 = str(row.get('三级模块', '')) if pd.notna(row.get('三级模块', '')) else ''
            function_name = str(row.get('功能过程', '')) if pd.notna(row.get('功能过程', '')) else ''
            subprocess_desc = str(row.get('子过程描述', '')) if pd.notna(row.get('子过程描述', '')) else ''

            # 记录一级模块和二级模块
            if module_l1 and module_l1 != 'nan' and module_l1.strip():
                module_l1_value = module_l1
                module_l1_set.add(module_l1)
            if module_l2 and module_l2 != 'nan' and module_l2.strip():
                module_l2_value = module_l2
                module_l2_set.add(module_l2)
            
            # 三级模块向下填充
            if module_l3 and module_l3 != 'nan' and module_l3.strip():
                module_l3_value = module_l3
            else:
                module_l3 = module_l3_value if module_l3_value else ''
            
            if module_l3 and module_l3.strip():
                module_l3_set.add(module_l3)
            
            # 功能过程向下填充
            if function_name and function_name != 'nan' and function_name.strip():
                function_name_value = function_name
            else:
                function_name = function_name_value if function_name_value else ''

            # 跳过空行
            if (not function_name or function_name == 'nan') and (not subprocess_desc or subprocess_desc == 'nan'):
                continue
            
            if not function_name or function_name == 'nan':
                continue
            
            # 只有有效数据才加入功能过程集合
            if function_name and function_name != 'nan':
                function_set.add(function_name)

            # 统计子过程
            if subprocess_desc and subprocess_desc != 'nan' and subprocess_desc.strip():
                sub_items = split_subprocess_description(subprocess_desc)
                subprocess_count += len(sub_items)

        # 生成输出文件名
        output_filename = f"{Path(filename).stem}_COSMIC.docx"
        word_path = OUTPUT_FOLDER / output_filename

        # 执行转换
        module_count = excel_to_word_conversion(excel_path, word_path)

        logger.info(f"转换完成：{output_filename}")

        return jsonify({
            'success': True,
            'message': '转换成功',
            'output_filename': output_filename,
            'module_count': module_count,
            'stats': {
                'l1_modules': len(module_l1_set),
                'l2_modules': len(module_l2_set),
                'l3_modules': len(module_l3_set),
                'functions': len(function_set),
                'subprocesses': subprocess_count
            }
        })

    except Exception as e:
        logger.error(f"转换失败：{e}", exc_info=True)
        return jsonify({'success': False, 'message': f'转换失败：{str(e)}'}), 500


@cosmic_bp.route('/download/<filename>', methods=['GET'])
def download_word(filename):
    """下载生成的 Word 文档"""
    try:
        word_path = OUTPUT_FOLDER / filename
        if not word_path.exists():
            return jsonify({'success': False, 'message': '文件不存在'}), 404

        return send_file(
            str(word_path),
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"下载失败：{e}", exc_info=True)
        return jsonify({'success': False, 'message': f'下载失败：{str(e)}'}), 500


@cosmic_bp.route('/stats', methods=['POST'])
def get_module_stats():
    """获取模块统计信息"""
    try:
        data = request.json
        filename = data.get('filename')

        if not filename:
            return jsonify({'success': False, 'message': '缺少文件名参数'}), 400

        excel_path = UPLOAD_FOLDER / filename
        if not excel_path.exists():
            return jsonify({'success': False, 'message': '文件不存在'}), 404

        # 读取数据
        df = read_excel_robust(excel_path)
        if df is None:
            return jsonify({'success': False, 'message': '无法读取 Excel 文件'}), 500

        # 统计各层级模块数量
        module_l1_set = set()
        module_l2_set = set()
        module_l3_set = set()
        function_set = set()
        total_subprocesses = 0
        
        module_l1_value = None
        module_l2_value = None
        module_l3_value = None
        function_name_value = None

        for idx, row in df.iterrows():
            module_l1 = str(row.get('一级模块', '')) if pd.notna(row.get('一级模块', '')) else ''
            module_l2 = str(row.get('二级模块', '')) if pd.notna(row.get('二级模块', '')) else ''
            module_l3 = str(row.get('三级模块', '')) if pd.notna(row.get('三级模块', '')) else ''
            function_name = str(row.get('功能过程', '')) if pd.notna(row.get('功能过程', '')) else ''
            subprocess_desc = str(row.get('子过程描述', '')) if pd.notna(row.get('子过程描述', '')) else ''

            # 记录一级模块和二级模块
            if module_l1 and module_l1 != 'nan' and module_l1.strip():
                module_l1_value = module_l1
                module_l1_set.add(module_l1)
            if module_l2 and module_l2 != 'nan' and module_l2.strip():
                module_l2_value = module_l2
                module_l2_set.add(module_l2)
            
            # 三级模块向下填充
            if module_l3 and module_l3 != 'nan' and module_l3.strip():
                module_l3_value = module_l3
            else:
                module_l3 = module_l3_value if module_l3_value else ''
            
            if module_l3 and module_l3.strip():
                module_l3_set.add(module_l3)
            
            # 功能过程向下填充
            if function_name and function_name != 'nan' and function_name.strip():
                function_name_value = function_name
            else:
                function_name = function_name_value if function_name_value else ''

            # 跳过空行
            if (not function_name or function_name == 'nan') and (not subprocess_desc or subprocess_desc == 'nan'):
                continue
            
            if not function_name or function_name == 'nan':
                continue
            
            # 只有有效数据才加入功能过程集合
            if function_name and function_name != 'nan':
                function_set.add(function_name)

            # 统计子过程数量
            if subprocess_desc and subprocess_desc != 'nan' and subprocess_desc.strip():
                sub_items = split_subprocess_description(subprocess_desc)
                total_subprocesses += len(sub_items)

        return jsonify({
            'success': True,
            'data': {
                'l1_count': len(module_l1_set),
                'l2_count': len(module_l2_set),
                'l3_count': len(module_l3_set),
                'function_count': len(function_set),
                'subprocess_count': total_subprocesses
            }
        })

    except Exception as e:
        logger.error(f"统计失败：{e}", exc_info=True)
        return jsonify({'success': False, 'message': f'统计失败：{str(e)}'}), 500


@cosmic_bp.route('/export-stats', methods=['POST'])
def export_module_stats():
    """导出模块统计信息为 Excel"""
    try:
        data = request.json
        filename = data.get('filename')

        if not filename:
            return jsonify({'success': False, 'message': '缺少文件名参数'}), 400

        excel_path = UPLOAD_FOLDER / filename
        if not excel_path.exists():
            return jsonify({'success': False, 'message': '文件不存在'}), 404

        # 读取数据
        df = read_excel_robust(excel_path)
        if df is None:
            return jsonify({'success': False, 'message': '无法读取 Excel 文件'}), 500

        # 重新执行转换逻辑获取详细的模块数据
        module_data = {}
        module_l1_value = None
        module_l2_value = None
        module_l3_value = None
        function_name_value = None
        
        # 统计用集合
        module_l1_set = set()
        module_l2_set = set()
        module_l3_set = set()
        function_set = set()
        total_subprocesses = 0

        for idx, row in df.iterrows():
            module_l1 = str(row.get('一级模块', '')) if pd.notna(row.get('一级模块', '')) else ''
            module_l2 = str(row.get('二级模块', '')) if pd.notna(row.get('二级模块', '')) else ''
            module_l3 = str(row.get('三级模块', '')) if pd.notna(row.get('三级模块', '')) else ''
            function_name = str(row.get('功能过程', '')) if pd.notna(row.get('功能过程', '')) else ''
            subprocess_desc = str(row.get('子过程描述', '')) if pd.notna(row.get('子过程描述', '')) else ''

            # 记录一级模块和二级模块
            if module_l1 and module_l1 != 'nan' and module_l1.strip():
                module_l1_value = module_l1
                module_l1_set.add(module_l1)
            if module_l2 and module_l2 != 'nan' and module_l2.strip():
                module_l2_value = module_l2
                module_l2_set.add(module_l2)
            
            # 三级模块向下填充
            if module_l3 and module_l3 != 'nan' and module_l3.strip():
                module_l3_value = module_l3
            else:
                module_l3 = module_l3_value if module_l3_value else ''
            
            if module_l3 and module_l3.strip():
                module_l3_set.add(module_l3)
            
            # 功能过程向下填充
            if function_name and function_name != 'nan' and function_name.strip():
                function_name_value = function_name
            else:
                function_name = function_name_value if function_name_value else ''

            # 跳过空行
            if (not function_name or function_name == 'nan') and (not subprocess_desc or subprocess_desc == 'nan'):
                continue
            
            if not function_name or function_name == 'nan':
                continue
            
            # 只有有效数据才加入功能过程集合
            if function_name and function_name != 'nan':
                function_set.add(function_name)

            # 按三级模块分组
            if module_l3 not in module_data:
                module_data[module_l3] = {
                    '一级模块': module_l1_value,
                    '二级模块': module_l2_value,
                    '三级模块': module_l3,
                    '功能点': {},
                    'subprocess_count': 0
                }

            if function_name not in module_data[module_l3]['功能点']:
                module_data[module_l3]['功能点'][function_name] = []

            # 收集子过程描述
            if subprocess_desc and subprocess_desc != 'nan' and subprocess_desc.strip():
                sub_items = split_subprocess_description(subprocess_desc)
                module_data[module_l3]['功能点'][function_name].extend(sub_items)
                module_data[module_l3]['subprocess_count'] += len(sub_items)
                total_subprocesses += len(sub_items)

        # 生成输出文件名
        output_filename = f"{Path(filename).stem}_模块统计.xlsx"
        output_path = OUTPUT_FOLDER / output_filename

        # 导出 Excel - 2个 Sheet
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Sheet 1: 汇总统计
            summary_data = [
                ['统计项', '数值'],
                ['一级模块数量', len(module_l1_set)],
                ['二级模块数量', len(module_l2_set)],
                ['三级模块数量', len(module_l3_set)],
                ['功能过程总数', len(function_set)],
                ['子过程总数', total_subprocesses]
            ]
            summary_df = pd.DataFrame(summary_data[1:], columns=summary_data[0])
            summary_df.to_excel(writer, sheet_name='汇总统计', index=False)
            
            # 调整汇总统计表列宽
            summary_ws = writer.sheets['汇总统计']
            summary_ws.column_dimensions['A'].width = 15
            summary_ws.column_dimensions['B'].width = 15
            
            # Sheet 2: 详细数据
            detail_data = []
            for module_l3, module_info in module_data.items():
                detail_data.append({
                    '一级模块名称': module_info['一级模块'],
                    '二级模块名称': module_info['二级模块'],
                    '三级模块名称': module_info['三级模块'],
                    '子过程数量': module_info['subprocess_count'],
                    'CFP总和': module_info['subprocess_count']  # CFP = 子过程数量
                })
            
            detail_df = pd.DataFrame(detail_data)
            detail_df.to_excel(writer, sheet_name='详细数据', index=False)
            
            # 调整详细数据表列宽
            detail_ws = writer.sheets['详细数据']
            detail_ws.column_dimensions['A'].width = 15
            detail_ws.column_dimensions['B'].width = 15
            detail_ws.column_dimensions['C'].width = 20
            detail_ws.column_dimensions['D'].width = 12
            detail_ws.column_dimensions['E'].width = 12

        logger.info(f"统计 Excel 已保存：{output_path}")

        # 直接返回文件下载
        return send_file(
            str(output_path),
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    except Exception as e:
        logger.error(f"导出统计失败：{e}", exc_info=True)
        return jsonify({'success': False, 'message': f'导出失败：{str(e)}'}), 500
