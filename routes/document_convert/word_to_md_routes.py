import os
from flask import Blueprint, request, render_template, send_file, abort, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from markdownify import markdownify as md
import io
import logging

logger = logging.getLogger(__name__)

word_to_md_bp = Blueprint('word_to_md', __name__, url_prefix='/api')

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def run_to_md(run):
    """将 Run 对象转换为 Markdown 格式"""
    text = run.text
    if not text:
        return ''
    
    # 处理下划线和删除线
    if run.underline:
        text = f"<u>{text}</u>"
    if run.strikethrough:
        text = f"~~{text}~~"
    
    # 处理加粗和斜体（需在其他格式之后）
    if run.bold and run.italic:
        text = f"***{text}***"
    elif run.bold:
        text = f"**{text}**"
    elif run.italic:
        text = f"*{text}*"
    
    return text


def detect_heading_level(paragraph):
    """
    检测段落的标题层级
    返回: (level, text) 或 None
    """
    style_name = paragraph.style.name if paragraph.style else ''
    style_lower = style_name.lower().strip()
    
    # 方法1: 标准 Heading 样式匹配（精确匹配）
    heading_map = {
        'heading 1': 1,
        'heading 2': 2,
        'heading 3': 3,
        'heading 4': 4,
        'heading 5': 5,
        'heading 6': 6,
        # 可能的变体
        'title': 1,
        'subtitle': 2,
        # 中文样式名称
        '标题': 1,
        '标题 1': 1,
        '标题 2': 2,
        '标题 3': 3,
        '标题 4': 4,
        '标题 5': 5,
        '标题 6': 6,
        '一级标题': 1,
        '二级标题': 2,
        '三级标题': 3,
        '四级标题': 4,
        '五级标题': 5,
        '六级标题': 6,
        'heading1': 1,
        'heading2': 2,
        'heading3': 3,
        'heading4': 4,
        'heading5': 5,
        'heading6': 6,
    }
    
    for key, level in heading_map.items():
        if style_lower == key or style_lower.startswith(key + ' '):
            text = ''.join(run_to_md(run) for run in paragraph.runs).strip()
            if text:
                logger.debug(f"检测到标准标题 [Level {level}]: {text[:50]}")
                return level, text
    
    # 方法2: 启发式检测 - 基于样式的潜在标题
    text_runs = [run for run in paragraph.runs if run.text.strip()]
    full_text = ''.join(run.text for run in text_runs).strip()
    
    if not full_text:
        return None
    
    # 检查是否为正文样式（排除已经处理的 Heading 样式）
    normal_styles = ['normal', 'normal0', 'body text', '', '正文', '正文文本', '默认段落字体']
    is_normal_style = style_lower in normal_styles
    
    # 方法3: 通过段落属性检测标题（间距、字体大小等）
    try:
        # 检查段落是否有特殊的间距属性（通常标题会有更大的段后间距）
        paragraph_format = paragraph.paragraph_format
        if paragraph_format:
            # 段后间距大于12磅可能是标题
            if paragraph_format.space_after and paragraph_format.space_after > 120:
                # 结合加粗判断
                if any(run.bold for run in text_runs):
                    text = ''.join(run_to_md(run) for run in paragraph.runs).strip()
                    if text and not text.endswith(('。', '.', '；', ';', '：', ':', '）', ')')):
                        level = 3  # 假设是三级标题
                        logger.debug(f"间距检测标题 [Level {level}]: {text[:50]}")
                        return level, text
    except:
        pass
    
    if not is_normal_style:
        return None
    
    # 启发式规则1: 短文本且全部加粗 → 可能是标题
    bold_runs = [run for run in text_runs if run.bold]
    bold_ratio = len(bold_runs) / len(text_runs) if text_runs else 0
    
    # 如果超过 80% 的 run 是加粗的，且文本较短，认为是标题
    if bold_ratio >= 0.8 and len(full_text) < 120:
        if not full_text.endswith(('。', '.', '；', ';', '：', ':', '）', ')', '？', '!', '！')):
            # 根据文本长度判断层级
            if len(full_text) < 15:
                level = 2  # ##
            elif len(full_text) < 30:
                level = 3  # ###
            elif len(full_text) < 50:
                level = 4  # ####
            elif len(full_text) < 80:
                level = 5  # #####
            else:
                level = 6  # ######
            
            logger.debug(f"启发式检测标题 [Level {level}, 加粗比例 {bold_ratio:.0%}]: {full_text[:50]}")
            return level, full_text
    
    # 启发式规则2: 单个 run 且加粗 → 很可能是标题
    if len(text_runs) == 1 and text_runs[0].bold and len(full_text) < 100:
        if not full_text.endswith(('。', '.', '；', ';', '：', ':', '？', '!', '！')):
            # 根据文本长度判断层级
            if len(full_text) < 20:
                level = 3  # ###
            elif len(full_text) < 50:
                level = 4  # ####
            else:
                level = 5  # #####
            logger.debug(f"单 run 加粗标题 [Level {level}]: {full_text[:50]}")
            return level, full_text
    
    # 启发式规则3: 以数字+顿号开头且短 → 可能是标题（如"1、背景"）
    if len(full_text) < 80 and (full_text.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')) or
                                 (full_text[0].isdigit() and len(full_text) > 1 and full_text[1] in ['、', '.', '．'])):
        if any(run.bold for run in text_runs):
            level = 4
            logger.debug(f"序号开头标题 [Level {level}]: {full_text[:50]}")
            return level, full_text
    
    return None


def para_to_md(paragraph):
    """将段落转换为 Markdown 格式"""
    # 首先尝试检测是否为标题
    heading_result = detect_heading_level(paragraph)
    if heading_result:
        level, text = heading_result
        return '#' * level + ' ' + text
    
    # 普通段落处理
    text = ''.join(run_to_md(run) for run in paragraph.runs).strip()
    return text if text else ''


def table_to_md(table):
    """将 Word 表格转换为 Markdown 表格格式"""
    try:
        md_lines = []
        rows = []
        
        # 提取所有行的数据
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # 清理单元格内容（去除多余空白）
                cell_text = cell.text.strip().replace('\n', ' ').replace('|', '\\|')
                cells.append(cell_text)
            rows.append(cells)
        
        if not rows:
            return ''
        
        # 确定列数（以第一行为准）
        num_cols = len(rows[0])
        
        # 生成表头
        header = '| ' + ' | '.join(rows[0]) + ' |'
        md_lines.append(header)
        
        # 生成分隔线
        separator = '| ' + ' | '.join(['---'] * num_cols) + ' |'
        md_lines.append(separator)
        
        # 生成数据行
        for row_data in rows[1:]:
            # 确保列数一致
            while len(row_data) < num_cols:
                row_data.append('')
            line = '| ' + ' | '.join(row_data[:num_cols]) + ' |'
            md_lines.append(line)
        
        logger.debug(f"转换表格: {len(rows)} 行 x {num_cols} 列")
        return '\n'.join(md_lines)
    except Exception as e:
        logger.warning(f"表格转换失败: {e}")
        return ''


def list_item_to_md(paragraph, is_numbered=False):
    """将列表项转换为 Markdown 格式"""
    text = ''.join(run_to_md(run) for run in paragraph.runs).strip()
    if not text:
        return ''
    
    prefix = '1. ' if is_numbered else '- '
    return prefix + text


def is_list_paragraph(paragraph):
    """检测段落是否为列表项"""
    # 方法1: 通过样式名称检测
    if paragraph.style and 'list' in paragraph.style.name.lower():
        return True
    
    # 方法2: 通过 XML 元素检测编号属性
    if paragraph._element.xpath('.//w:numPr'):
        return True
    
    # 方法3: 通过段落格式属性检测
    try:
        if hasattr(paragraph, 'style') and paragraph.style:
            style_id = paragraph.style.id
            if style_id and ('list' in style_id.lower() or 'num' in style_id.lower()):
                return True
    except:
        pass
    
    return False


def get_list_level(paragraph):
    """获取列表项的层级（从0开始）"""
    try:
        # 尝试从段落的编号属性中获取层级
        num_pr = paragraph._element.xpath('.//w:numPr')
        if num_pr:
            # 查找 w:ilvl 元素
            ilvl = num_pr[0].xpath('.//w:ilvl')
            if ilvl:
                # 获取层级值
                val = ilvl[0].xpath('.//w:val')
                if val:
                    return int(val[0].text) if val[0].text else 0
        return 0
    except:
        return 0


def is_numbered_list(paragraph):
    """检测是否为有序列表"""
    try:
        num_pr = paragraph._element.xpath('.//w:numPr')
        if num_pr:
            # 查找编号格式类型
            num_fmt = num_pr[0].xpath('.//w:numFmt')
            if num_fmt:
                val = num_fmt[0].xpath('.//w:val')
                if val and val[0].text:
                    # 如果是数字格式，则为有序列表
                    return val[0].text in ['decimal', 'lowerRoman', 'upperRoman', 'lowerLetter', 'upperLetter']
            # 默认认为有编号属性的是有序列表
            return True
        return False
    except:
        return False


def docx_to_markdown(docx_path):
    """将 Word 文档转换为 Markdown 格式"""
    logger.info(f"开始转换 Word 文档: {docx_path}")
    
    doc = Document(docx_path)
    md_lines = []
    
    # 统计信息
    para_count = 0
    table_count = 0
    heading_count = 0
    list_count = 0
    
    # 追踪列表状态
    in_list = False
    
    # 处理段落
    for para in doc.paragraphs:
        # 检查是否为列表项
        if is_list_paragraph(para):
            # 获取列表层级和类型
            level = get_list_level(para)
            is_numbered = is_numbered_list(para)
            text = ''.join(run_to_md(run) for run in para.runs).strip()
            
            if text:
                # 添加缩进（每级2个空格）
                indent = '  ' * level
                prefix = '1. ' if is_numbered else '- '
                md_line = indent + prefix + text
                md_lines.append(md_line)
                list_count += 1
                in_list = True
                continue
        
        # 普通段落或标题处理
        md_line = para_to_md(para)
        if md_line:
            # 如果之前是列表，现在不是，添加空行分隔
            if in_list:
                md_lines.append('')
                in_list = False
            
            md_lines.append(md_line)
            para_count += 1
            
            # 统计标题数量
            if md_line.startswith('#'):
                heading_count += 1
    
    # 处理表格
    for table in doc.tables:
        table_md = table_to_md(table)
        if table_md:
            # 在表格前后添加空行
            if md_lines and md_lines[-1]:
                md_lines.append('')
            md_lines.append(table_md)
            md_lines.append('')
            table_count += 1
    
    # 处理段落中的换行符（保留合理的换行）
    final_lines = []
    for line in md_lines:
        # 将多个连续换行合并为两个（Markdown段落分隔）
        if line.strip():
            final_lines.append(line)
        elif final_lines and final_lines[-1]:
            # 只在非空行后添加一个空行
            final_lines.append('')
    
    logger.info(f"转换完成: {para_count} 个段落, {table_count} 个表格, {heading_count} 个标题, {list_count} 个列表项")
    
    return '\n'.join(final_lines)


# ============================================================================
# Vue 前端 API 接口（返回 JSON）
# ============================================================================

@word_to_md_bp.route('/word-to-md/convert', methods=['POST'])
def word_to_md_convert_api():
    """Word 转 Markdown API（Vue 前端使用，返回 JSON）"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '请选择有效的 Word 文件（.docx）'}), 400
        
        # 检查文件扩展名
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': '请选择有效的 Word 文件（.docx）'}), 400

        # 获取原始文件名（保留中文）
        original_filename = file.filename
        
        # 安全检查：防止路径遍历攻击
        safe_filename = os.path.basename(original_filename)
        
        # 生成 MD 文件名
        filename_without_ext = os.path.splitext(safe_filename)[0]
        md_filename = filename_without_ext + '.md'
        
        # 保存文件
        filepath = os.path.join(UPLOAD_FOLDER, safe_filename)
        file.save(filepath)

        try:
            markdown_text = docx_to_markdown(filepath)
        except Exception as e:
            logger.error(f"转换失败: {e}", exc_info=True)
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': f'转换失败：{str(e)}'}), 500

        # 删除临时上传文件
        if os.path.exists(filepath):
            os.remove(filepath)

        return jsonify({
            'success': True,
            'message': '转换成功',
            'markdown': markdown_text,
            'filename': md_filename
        })
    
    except Exception as e:
        logger.error(f"服务器错误: {e}", exc_info=True)
        return jsonify({'success': False, 'message': f'服务器错误：{str(e)}'}), 500


@word_to_md_bp.route('/word-to-md/debug-styles', methods=['POST'])
def debug_word_styles():
    """调试接口：分析 Word 文档的样式结构（帮助诊断转换问题）"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '请选择有效的 Word 文件（.docx）'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': '请选择有效的 Word 文件（.docx）'}), 400

        safe_filename = os.path.basename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, safe_filename)
        file.save(filepath)

        try:
            doc = Document(filepath)
            style_analysis = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'paragraphs': [],
                'unique_styles': set()
            }
            
            # 分析每个段落的样式
            for idx, para in enumerate(doc.paragraphs[:50]):  # 只分析前50个段落
                text_preview = para.text[:80] if para.text else '(空)'
                style_name = para.style.name if para.style else '(无样式)'
                style_analysis['unique_styles'].add(style_name)
                
                # 检查格式特征
                is_bold = any(run.bold for run in para.runs if run.text.strip())
                is_italic = any(run.italic for run in para.runs if run.text.strip())
                has_runs = len(para.runs)
                
                style_analysis['paragraphs'].append({
                    'index': idx,
                    'text': text_preview,
                    'style': style_name,
                    'is_bold': is_bold,
                    'is_italic': is_italic,
                    'run_count': has_runs,
                    'length': len(para.text)
                })
            
            # 转换为可序列化的格式
            style_analysis['unique_styles'] = sorted(list(style_analysis['unique_styles']))
            
            logger.info(f"样式分析完成: {len(style_analysis['unique_styles'])} 种样式")
            
            return jsonify({
                'success': True,
                'analysis': style_analysis
            })
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)
    
    except Exception as e:
        logger.error(f"样式分析失败: {e}", exc_info=True)
        return jsonify({'success': False, 'message': f'分析失败：{str(e)}'}), 500


# ============================================================================
# 传统 HTML 模板接口 (已废弃，使用 Vue 页面)
# 保留 POST 用于文件上传处理
# ============================================================================

@word_to_md_bp.route('/word-to-md', methods=['POST'])
def word_to_md():
    if request.method == 'POST':
        if 'word_file' not in request.files:
            return render_template('word_to_md.html', error='没有选择文件')
        file = request.files['word_file']
        if file.filename == '':
            return render_template('word_to_md.html', error='请选择有效的 Word 文件（.docx）')
        
        # 检查文件扩展名
        if not allowed_file(file.filename):
            return render_template('word_to_md.html', error='请选择有效的 Word 文件（.docx）')

        # 获取原始文件名（不进行 secure_filename 处理，保留中文）
        original_filename = file.filename
        
        # 安全检查：确保文件名不包含路径遍历攻击
        # 只取文件名部分，防止 ../../etc/passwd 这种攻击
        import os
        safe_filename = os.path.basename(original_filename)
        
        # 记录文件名处理过程
        from flask import current_app
        logger = current_app.logger if hasattr(current_app, 'logger') else None
        if logger:
            logger.info(f"[WORD_TO_MD] 原始文件名：{original_filename}")
            logger.info(f"[WORD_TO_MD] 安全处理后：{safe_filename}")
        
        # 生成输出文件名：保持原名，只改变扩展名
        filename_without_ext = os.path.splitext(safe_filename)[0]
        md_filename = filename_without_ext + '.md'
        
        if logger:
            logger.info(f"[WORD_TO_MD] 生成的 MD 文件名：{md_filename}")
        
        filepath = os.path.join(UPLOAD_FOLDER, safe_filename)
        file.save(filepath)

        try:
            markdown_text = docx_to_markdown(filepath)
        except Exception as e:
            if os.path.exists(filepath):
                os.remove(filepath)
            return render_template('word_to_md.html', error=f'转换失败：{str(e)}')

        # 删除上传的文件
        if os.path.exists(filepath):
            os.remove(filepath)

        # 将 markdown 文本写入内存文件，返回下载
        md_bytes = markdown_text.encode('utf-8')
        md_file = io.BytesIO(md_bytes)
        md_file.seek(0)

        return send_file(
            md_file,
            as_attachment=True,
            download_name=md_filename,
            mimetype='text/markdown'
        )

    return jsonify({'success': False, 'message': '请使用 Vue 前端页面访问'}), 404
